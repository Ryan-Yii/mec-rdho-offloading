import sys
from pathlib import Path

from docx import Document

from tools.revise_round2 import set_rich_para


def test_set_rich_para_writes_math_segments_as_inline_omml():
    doc = Document()
    paragraph = doc.add_paragraph()
    set_rich_para(paragraph, [('Let ', False), (r'x_k^{\mathrm{new}}=x_k+1', True), ('.', False)])

    math_nodes = paragraph._p.xpath('.//*[local-name()="oMath"]')
    assert len(math_nodes) == 1
    assert 'Let ' in paragraph.text


def test_replace_terms_preserves_inline_picture(tmp_path):
    from PIL import Image
    from docx.shared import Inches
    from tools.revise_round2 import replace_terms_preserving_drawings

    image_path = tmp_path / 'tiny.png'
    Image.new('RGB', (10, 10), 'white').save(image_path)
    doc = Document()
    paragraph = doc.add_paragraph('resource ratio ')
    paragraph.add_run().add_picture(str(image_path), width=Inches(0.1))

    replace_terms_preserving_drawings(doc, {'resource ratio': 'computation-control value'})

    assert 'computation-control value' in paragraph.text
    assert len(paragraph._p.xpath('.//*[local-name()="drawing"]')) == 1
