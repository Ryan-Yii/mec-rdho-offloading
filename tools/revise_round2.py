from __future__ import annotations

import copy
import json
import shutil
import tempfile
import zipfile
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from lxml import etree
from scipy.stats import rankdata, wilcoxon

import os

REPO = Path(__file__).resolve().parents[1]
SOURCE = Path(os.environ.get('RDHO_SOURCE_DOCX', REPO / 'source_marked.docx'))
INTERMEDIATE = Path(os.environ.get('RDHO_INTERMEDIATE_DOCX', REPO / 'RDHO_round2_intermediate.docx'))
OUTPUT = Path(os.environ.get('RDHO_OUTPUT_DOCX', REPO / 'RDHO_SCI_round2_marked.docx'))

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
W14_NS = 'http://schemas.microsoft.com/office/word/2010/wordml'
W15_NS = 'http://schemas.microsoft.com/office/word/2012/wordml'
NS = {'w': W_NS, 'w14': W14_NS, 'w15': W15_NS}


def comment_ids(paragraph):
    ids=[]
    for e in paragraph._p.xpath('.//*[local-name()="commentRangeStart" or local-name()="commentReference"]'):
        cid=e.get(qn('w:id'))
        if cid is not None and cid not in ids: ids.append(cid)
    return ids


def clear_preserve_ppr(paragraph):
    for child in list(paragraph._p):
        if child.tag != qn('w:pPr'):
            paragraph._p.remove(child)


def add_comment_start(paragraph, cid):
    e=OxmlElement('w:commentRangeStart'); e.set(qn('w:id'),cid); paragraph._p.append(e)


def add_comment_end(paragraph,cid):
    e=OxmlElement('w:commentRangeEnd'); e.set(qn('w:id'),cid); paragraph._p.append(e)
    r=OxmlElement('w:r'); rpr=OxmlElement('w:rPr'); sty=OxmlElement('w:rStyle'); sty.set(qn('w:val'),'CommentReference'); rpr.append(sty); r.append(rpr)
    ref=OxmlElement('w:commentReference'); ref.set(qn('w:id'),cid); r.append(ref); paragraph._p.append(r)


def format_run(run,size=12,bold=None,italic=None,highlight=True,font='Times New Roman'):
    run.font.name=font
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'),font)
    run.font.size=Pt(size)
    if bold is not None: run.bold=bold
    if italic is not None: run.italic=italic
    if highlight: run.font.highlight_color=WD_COLOR_INDEX.YELLOW


def set_para(paragraph,text,size=12,bold=None,alignment=None,highlight=True):
    ids=comment_ids(paragraph); clear_preserve_ppr(paragraph)
    for cid in ids: add_comment_start(paragraph,cid)
    if text:
        r=paragraph.add_run(text); format_run(r,size,bold=bold,highlight=highlight)
    for cid in reversed(ids): add_comment_end(paragraph,cid)
    if alignment is not None: paragraph.alignment=alignment


def set_rich_para(paragraph,segments,size=12,alignment=None):
    """segments: [(text, is_math)]; math text is LaTeX and becomes inline OMML."""
    ids=comment_ids(paragraph); clear_preserve_ppr(paragraph)
    for cid in ids: add_comment_start(paragraph,cid)
    for text,is_math in segments:
        if is_math:
            paragraph._p.append(build_math(text))
        else:
            r=paragraph.add_run(text)
            format_run(r,size,highlight=True,font='Times New Roman')
    for cid in reversed(ids): add_comment_end(paragraph,cid)
    if alignment is not None: paragraph.alignment=alignment


def build_math(formula:str):
    import subprocess
    with tempfile.TemporaryDirectory() as td:
        td=Path(td); md=td/'e.md'; out=td/'e.docx'
        md.write_text(f'$$\n{formula}\n$$',encoding='utf-8')
        subprocess.run(['pandoc',str(md),'-o',str(out)],check=True)
        d=Document(out)
        node=d.paragraphs[0]._p.xpath('.//*[local-name()="oMath"]')[0]
        node=copy.deepcopy(node)
        # yellow highlight each math run
        for mr in node.xpath('.//*[local-name()="r" and namespace-uri()="http://schemas.openxmlformats.org/officeDocument/2006/math"]'):
            rpr=next((c for c in mr if c.tag==qn('m:rPr')),None)
            if rpr is None:
                rpr=OxmlElement('m:rPr'); mr.insert(0,rpr)
            wrpr=next((c for c in rpr if c.tag==qn('w:rPr')),None)
            if wrpr is None:
                wrpr=OxmlElement('w:rPr'); rpr.append(wrpr)
            hi=OxmlElement('w:highlight'); hi.set(qn('w:val'),'yellow'); wrpr.append(hi)
        return node


def set_equation(paragraph,formula,number):
    ids=comment_ids(paragraph); clear_preserve_ppr(paragraph)
    # center + right tab
    ppr=paragraph._p.get_or_add_pPr()
    tabs=OxmlElement('w:tabs')
    for val,pos in [('center','4860'),('right','9720')]:
        t=OxmlElement('w:tab'); t.set(qn('w:val'),val); t.set(qn('w:pos'),pos); tabs.append(t)
    ppr.append(tabs)
    for cid in ids: add_comment_start(paragraph,cid)
    r=OxmlElement('w:r'); r.append(OxmlElement('w:tab')); paragraph._p.append(r)
    paragraph._p.append(build_math(formula))
    paragraph.add_run('\t')
    rr=paragraph.add_run(f'({number})'); format_run(rr,11,highlight=True)
    for cid in reversed(ids): add_comment_end(paragraph,cid)


def set_cell(cell,text,header=False,size=9):
    cell.text=''; p=cell.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(str(text)); format_run(r,size,bold=header,highlight=True)
    p.paragraph_format.space_before=Pt(0); p.paragraph_format.space_after=Pt(0); p.paragraph_format.line_spacing=1.0
    cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER


def set_border(cell,**edges):
    tcPr=cell._tc.get_or_add_tcPr(); borders=tcPr.first_child_found_in('w:tcBorders')
    if borders is None: borders=OxmlElement('w:tcBorders'); tcPr.append(borders)
    for edge,data in edges.items():
        e=borders.find(qn('w:'+edge))
        if e is None: e=OxmlElement('w:'+edge); borders.append(e)
        for k,v in data.items(): e.set(qn('w:'+k),str(v))


def replace_terms_preserving_drawings(doc, replacements):
    """Replace text only in textual runs; never rewrite runs containing drawings."""
    paragraphs = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    for paragraph in paragraphs:
        for run in paragraph.runs:
            if run._r.xpath('.//*[local-name()="drawing" or local-name()="pict"]'):
                continue
            original = run.text
            updated = original
            for old, new in replacements.items():
                updated = updated.replace(old, new)
            if updated != original:
                run.text = updated


def set_table(table,data,widths,font=8.5):
    while len(table.rows)<len(data): table.add_row()
    while len(table.columns)<max(map(len,data)): table.add_column(Inches(.5))
    table.autofit=False
    if widths:
        grid = table._tbl.tblGrid
        gridcols = list(grid.gridCol_lst)
        for i,w in enumerate(widths):
            table.columns[i].width = Inches(w)
            if i < len(gridcols):
                gridcols[i].set(qn('w:w'), str(int(Inches(w).emu / 635)))
            for c in table.columns[i].cells:
                c.width=Inches(w)
                tcw = c._tc.get_or_add_tcPr().first_child_found_in('w:tcW')
                if tcw is None:
                    tcw=OxmlElement('w:tcW'); c._tc.get_or_add_tcPr().append(tcw)
                tcw.set(qn('w:w'), str(int(Inches(w).emu / 635)))
                tcw.set(qn('w:type'),'dxa')
    for i,row in enumerate(data):
        for j in range(len(table.columns)): set_cell(table.cell(i,j),row[j] if j<len(row) else '',header=(i==0),size=font)
    for i in range(len(data),len(table.rows)):
        for j in range(len(table.columns)): set_cell(table.cell(i,j),'',size=font)
    nil={'val':'nil'}; single={'val':'single','sz':'8','space':'0','color':'000000'}
    for row in table.rows:
        for c in row.cells: set_border(c,top=nil,bottom=nil,left=nil,right=nil,insideH=nil,insideV=nil)
    for c in table.rows[0].cells: set_border(c,top=single,bottom=single)
    for c in table.rows[len(data)-1].cells: set_border(c,bottom=single)


def fmt(m,s,d=3): return f'{m:.{d}f} ± {s:.{d}f}'


def summary(df,col,key): return df[df[col]==key].iloc[0]


def replace_picture(paragraph,path,width):
    clear_preserve_ppr(paragraph); paragraph.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run().add_picture(str(path),width=Inches(width))


def ablation_stats(raw):
    piv=raw.pivot(index='run_id',columns='algorithm',values='fitness')
    full=piv['RDHO-full'].to_numpy(float)
    names=[c for c in piv.columns if c!='RDHO-full']
    rec=[]
    for name in names:
        x=piv[name].to_numpy(float)
        st=wilcoxon(full,x,alternative='two-sided',zero_method='wilcox')
        diff=full-x
        nz=diff[np.abs(diff)>1e-12]
        if len(nz):
            ranks=rankdata(np.abs(nz)); pos=ranks[nz>0].sum(); neg=ranks[nz<0].sum(); rb=(pos-neg)/(pos+neg)
        else: rb=0
        rec.append({'variant':name,'W':float(st.statistic),'p':float(st.pvalue),'median':float(np.median(diff)),'rank_biserial':float(rb)})
    order=np.argsort([r['p'] for r in rec]); run=0
    for rank,idx in enumerate(order):
        val=min(1,(len(rec)-rank)*rec[idx]['p']); run=max(run,val); rec[idx]['p_holm']=run
    return pd.DataFrame(rec)


def patch_comments(docx_path:Path,out_path:Path):
    updates={
      '1':'已按IJSNet审阅稿方向统一字体、表格与章节结构，并保留作者信息、批注和高亮供导师审阅；正式投稿时需另存匿名版本并清除个人元数据。',
      '3':'已将标题统一为“RIME-DBO-Based QoE- and Fairness-Aware Task Offloading in Mobile Edge Computing”，并在贡献与实验结论中限定为完整RDHO-full求解流程，避免将优势过度归因于单一融合算子。',
      '5':'回复0712：已采纳简化建议，标题采用“RIME-DBO-Based QoE- and Fairness-Aware Task Offloading in Mobile Edge Computing”，突出方法来源与研究问题。',
      '7':'回复0712：摘要已按“背景与意义—研究挑战—本文方法—主要结论与边界”的顺序重写。',
      '9':'回复0712：摘要不再堆叠具体表格数值，仅保留研究设计、验证范围和保守结论。',
      '11':'已统一为“five-metric scalarised objective”或“weighted multi-metric formulation”。本文使用单一加权标量目标，不再误称为Pareto型multi-objective算法。',
      '13':'回复0712：关键词中已删除“hybrid”，并统一为task offloading、computation control、QoE和priority-aware fairness等准确术语。',
      '15':'已按倒三角结构扩写Introduction：应用背景、MEC调度挑战、现有研究缺口、本文贡献和组织结构。',
      '17':'回复0712：已使用原创云—边—端架构图，并在正文明确固定接入边缘、云转发路径和本地/边缘/云三种执行方式，未复用会议论文原图。',
      '19':'回复0712：正文引用统一为[x]序号形式，参考文献按首次出现顺序排列。',
      '21':'回复0712：Related Work已扩充学习式卸载、AoI、QoE/公平性和混合元启发式研究，并明确本文与前期第四章工作的扩展关系。',
      '23':'已按高被引综述常见框架分类介绍energy-delay、learning-based、freshness-aware、QoE/fairness和metaheuristic工作。',
      '25':'回复0712：已补全云—边—端系统、任务映射、时延与设备侧能耗模型；AoI现准确表述为single-epoch surrogate，不再声称峰值AoI或队列级AoI。',
      '27':'回复0712：核心服务率、时延、能耗、AoI与适应度公式均单行陈列并右对齐编号；分段式仅用于同一物理量在三种执行模式下的统一定义。',
      '29':'已精简Table 1并统一符号体系：集合使用花体字母，集合大小使用普通大写字母，新增全局任务平均QoE定义；少量紧密关联符号按语义分组以控制版面。',
      '31':'回复0712：Table 1仅保留核心符号，并补充集合/数量、全局平均QoE、搜索适应度和报告适应度；其余任务参数在正文与参数表中定义。',
      '33':'已完成。全部编号公式采用居中公式、右侧编号的统一格式。',
      '35':'已完成。约束与评价语义合并为连续正文，避免过度细分小节。',
      '37':'已完成。新增与修改的公式编号和交叉引用均保持一致。',
      '39':'已完成。Algorithm 1采用Require/Ensure与编号步骤，并明确搜索适应度、报告适应度、边界解码和可选局部细化。',
      '41':'回复0712：Algorithm 1已按三线表形式重构；算法段中的随机变量、上下界、个体下标和概率表达均统一为专业数学格式。',
      '43':'回复0712：全文表格采用统一三线表、Times New Roman、无竖线、固定列宽与一致的对齐/间距；Table 6已重新排版。',
      '45':'已完成。参考文献从[1]开始并按正文首次引用顺序排列，正文编号同步。',
      '47':'已补充2024—2025年MEC/offloading综述与相关工作；实验部分增加30次配对统计、消融、规模和敏感性分析，并对贡献归因保持保守。'
    }
    with tempfile.TemporaryDirectory() as td:
        td=Path(td)
        with zipfile.ZipFile(docx_path) as z: z.extractall(td)
        cpath=td/'word/comments.xml'; expath=td/'word/commentsExtended.xml'
        root=etree.parse(str(cpath)); comments=root.getroot()
        infos=[]
        for c in comments.findall('w:comment',NS):
            cid=c.get(f'{{{W_NS}}}id'); author=c.get(f'{{{W_NS}}}author','')
            paras=c.findall('w:p',NS); paraid=paras[0].get(f'{{{W14_NS}}}paraId') if paras else None
            infos.append((int(cid),cid,author,paraid))
            if cid in updates:
                ts=c.findall('.//w:t',NS)
                if not ts:
                    p=paras[0] if paras else etree.SubElement(c,f'{{{W_NS}}}p')
                    r=etree.SubElement(p,f'{{{W_NS}}}r'); t=etree.SubElement(r,f'{{{W_NS}}}t'); ts=[t]
                ts[0].text=updates[cid]
                for t in ts[1:]: t.text=''
        root.write(str(cpath),xml_declaration=True,encoding='UTF-8',standalone=True)

        ex=etree.parse(str(expath)); exroot=ex.getroot(); bypara={e.get(f'{{{W15_NS}}}paraId'):e for e in exroot.findall('w15:commentEx',NS)}
        paraid_by_cid={cid:paraid for _,cid,_,paraid in infos if paraid}
        for numeric_id,cid,author,paraid in sorted(infos):
            if not paraid or paraid not in bypara: continue
            e=bypara[paraid]
            parent_attr=f'{{{W15_NS}}}paraIdParent'
            if author=='webuser':
                e.attrib.pop(parent_attr,None)
            elif author=='祎宝':
                # The review file uses paired comment IDs: teacher 0, author reply 1, etc.
                parent_paraid=paraid_by_cid.get(str(numeric_id-1))
                if parent_paraid:
                    e.set(parent_attr,parent_paraid)
        ex.write(str(expath),xml_declaration=True,encoding='UTF-8',standalone=True)
        with zipfile.ZipFile(out_path,'w',zipfile.ZIP_DEFLATED) as z:
            for f in td.rglob('*'):
                if f.is_file(): z.write(f,f.relative_to(td))


def main():
    main_summary=pd.read_csv(REPO/'results/summary/main_30_summary_mean_std.csv')
    stats=pd.read_csv(REPO/'results/summary/wilcoxon_fitness_results.csv')
    ab_summary=pd.read_csv(REPO/'results/summary/ablation_30_summary_mean_std.csv')
    ab_raw=pd.read_csv(REPO/'results/raw/ablation_30_raw_results.csv')
    scale=pd.read_csv(REPO/'results/summary/scalability_summary_mean_std.csv')
    weight=pd.read_csv(REPO/'results/sensitivity/summary/weight_sensitivity_summary_mean_std.csv')
    penalty=pd.read_csv(REPO/'results/sensitivity/summary/dynamic_penalty_sensitivity_summary_mean_std.csv')
    abs_stats=ablation_stats(ab_raw)
    abs_stats.to_csv(REPO/'results/summary/ablation_wilcoxon_results.csv',index=False)

    doc=Document(SOURCE); p=doc.paragraphs
    rdho=summary(main_summary,'algorithm','RDHO')
    baselines=['RIME','DBO','TLBO-HHO','CWTSSA','Greedy-ED']
    imps={b:100*(summary(main_summary,'algorithm',b).fitness_mean-rdho.fitness_mean)/summary(main_summary,'algorithm',b).fitness_mean for b in baselines}
    imptext=', '.join(f'{imps[b]:.1f}% versus {b}' for b in baselines)

    # Core positioning and attribution.
    set_para(p[6],('Mobile edge computing (MEC) brings computation closer to resource-constrained devices, but heterogeneous workloads require scheduling decisions that balance technical efficiency, information freshness, model-based service utility, and equitable user outcomes. '
        'This paper develops RDHO, an end-to-end solver whose population stage uses RIME- and DBO-inspired operators and whose complete RDHO-full configuration also includes greedy seeding and coordinate refinement. Each task selects local, edge, or cloud execution together with a bounded normalised computation-control value. '
        'The configured scalar objective combines device-side energy, delay, a single-epoch AoI surrogate, a priority-weighted QoE proxy, and priority-aware per-user utility fairness. Search and reporting fitness are separated to maintain a common final evaluation scale. Thirty paired simulations, corrected non-parametric tests, component analysis, scalability tests, and sensitivity studies show that the complete RDHO-full procedure obtains the lowest reporting objective in the implemented suite. This evidence concerns the configured end-to-end solver and does not isolate superiority of the RIME-DBO population operator from its seeding and refinement stages.'),12)
    set_para(p[13],('The main contributions are fourfold. First, we formulate a reproducible cloud-edge-device offloading abstraction with task-level execution modes and bounded normalised computation-control values, together with explicit limitations on CPU-capacity interpretation. Second, we develop the complete RDHO solver, whose population stage uses role-conditioned RIME- and DBO-inspired updates and whose full configuration includes greedy seeding and coordinate refinement; these internal mechanisms are analysed without assuming that each is independently dominant. Third, we separate iteration-dependent search fitness from fixed-reference reporting fitness and evaluate parents and candidates under the same current penalty coefficient. Fourth, we provide paired 30-run comparisons, two-sided corrected statistics, paired component analysis, scalability and sensitivity studies, and a complete reproducibility artefact.'),12)
    set_para(p[14],f'Under the configured end-to-end procedures, RDHO-full obtains a mean fixed-reference reporting fitness of {rdho.fitness_mean:.4f} and reduces that objective by {imptext}. These comparisons support the complete RDHO-full procedure, not a strictly NFE-matched superiority claim for the population-level RIME-DBO fusion. Raw metrics, runtime, and NFE are reported separately.',12)
    set_para(p[23],('The resulting gap is a reproducible user-oriented scheduler that jointly reports technical efficiency, a freshness surrogate, model-based QoE, and priority-aware per-user utility fairness while clearly separating search guidance from final evaluation. The present study evaluates the complete RDHO pipeline; it does not claim that the RIME-DBO operator alone has been isolated from greedy initialisation and final coordinate refinement.'),12)

    # Symbols and global QoE.
    set_rich_para(p[26],[('The considered system uses a three-tier cloud-edge-device architecture. Let ',False),(r'\mathcal M,\mathcal G,\mathcal C,\mathcal N',True),(' denote the sets of mobile devices, access-edge servers, cloud servers, and tasks. Their cardinalities are ',False),(r'M=\mathrm{card}(\mathcal M),\;G=\mathrm{card}(\mathcal G),\;C=\mathrm{card}(\mathcal C),\;N=\mathrm{card}(\mathcal N)',True),('. Each device has a fixed access edge during one scheduling epoch, and each edge has a fixed cloud forwarding target.',False)],12)
    set_rich_para(p[83],[('The global task-average QoE used in the scalar objective is ',False),(r'\bar q=\frac{1}{N}\sum_{i\in\mathcal N}q_i',True),('. Thus, users generating more tasks contribute more terms to q̄, whereas fairness J first aggregates tasks by active user so that each active user receives one fairness entry. In the order energy, delay, AoI, QoE, and fairness, the default weights are (0.15, 0.15, 0.20, 0.25, 0.25) and sum to one. These are engineering preference coefficients rather than statistically independent contributions.',False)],12)
    set_rich_para(p[108],[('The model is a nonconvex mixed discrete-continuous optimisation problem. Even before optimising computation-control values, ',False),(r'N=\mathrm{card}(\mathcal N)',True),(' tasks produce ',False),(r'3^N',True),(' possible execution-mode assignments.',False)],12)

    # Model coefficients: rationale and limitations.
    set_para(p[43],('The load terms in Eqs. (1)-(3) count tasks using the corresponding device, edge, and cloud nodes. The affine control factors (0.35,0.65), (0.40,0.60), and (0.45,0.55) preserve positive service rates and monotonic response to rᵢ, while the congestion factors 0.07, 0.035, and 0.020 impose progressively weaker per-task degradation at device, edge, and cloud layers. These dimensionless values are inherited from the preliminary thesis simulator and were fixed before algorithm comparison; they were not fitted to physical processor traces or tuned separately for any algorithm. The service overheads δᴱ=0.010 s and δᶜ=0.055 s are likewise fixed simulation parameters.'),12)
    set_para(p[50],('The coefficients 0.015 J/s and 0.010 J/s are fixed simulation-level device-active overhead proxies for edge and cloud service. They are inherited from the preliminary simulator, applied identically to every algorithm, and are not claimed as hardware-calibrated power measurements. Powered edge/cloud infrastructure energy is not modelled; the reported energy metric is therefore a device-side proxy. Sensitivity of algorithm ranking to these service-model coefficients remains outside the present evaluation and is stated as a limitation.'),12)

    # Algorithm description and corrected foraging bounds.
    set_para(p[112],('RDHO addresses the five-metric problem in Section 3 through a complete solver pipeline. Its population stage uses RIME-inspired exploration and DBO-inspired role-conditioned movements, while RDHO-full additionally uses greedy seeding and final coordinate refinement. RDHO-core denotes the population-search stage without final refinement. The experiments evaluate these configurations but do not isolate a claim that the RIME-DBO operator alone is superior to equivalently seeded and refined parent algorithms.'),12)
    set_para(p[115],('The population update is role-conditioned: producers use a weighted combination of RIME- and DBO-inspired candidate coordinates, followers alternate between best-guided puncture and bound-aware foraging, and scouts apply theft or Cauchy perturbations. These choices define the configured RDHO search architecture; their individual necessity is assessed empirically rather than assumed from the hybrid label.'),12)
    set_para(p[117],('Each individual is an N×2 matrix containing (zᵢ,rᵢ) for every task. Metaheuristic updates use a continuous relaxation of both coordinates. Before evaluation, zᵢ is decoded by nearest-integer projection and clipping onto {0,1,2}, while rᵢ is clipped to [0.2,1.0]. This encoding imposes an algorithmic neighbourhood on the categorical modes—for example, labels 0 and 1 are closer in the relaxed coordinate than labels 0 and 2—even though that neighbourhood does not represent a physical distance among local, edge, and cloud execution.'),12)
    set_rich_para(p[121],[('The initial population contains equal Gaussian and uniform subsets. For the Gaussian subset, ',False),(r'z\sim\mathcal N(1,0.7^2)',True),(' and ',False),(r'r\sim\mathcal N(0.68,0.18^2)',True),('; for the uniform subset, ',False),(r'z\sim\mathcal U(0,2)',True),(' and ',False),(r'r\sim\mathcal U(0.2,1)',True),('. The first individual is a greedy coordinate seed, and up to three additional seeds use zero-mean Gaussian perturbations with standard deviation 0.08 before decoding and clipping.',False)],12)
    set_rich_para(p[122],[('Let ',False),(r'p=t/T_{\max}',True),(' and let ',False),(r'\mathcal D(t)',True),(' be the mean, over all relaxed decision dimensions, of the population standard deviation. The producer ratio is ',False),(r'\mathrm{clip}(0.28-0.10p+0.08\mathcal D(t),0.14,0.34)',True),(', the scout ratio is ',False),(r'\mathrm{clip}(0.08+0.08\mathcal D(t)+0.04p,0.08,0.20)',True),(', and the follower ratio receives the remainder subject to a minimum of 0.40. Integer role counts are obtained by truncation with at least one member per role; the top 10% are preserved as elites.',False)],12)
    set_equation(p[125],r'X_k^{\mathrm{new}}=w(t)X_k^{\mathrm{RIME}}+\left[1-w(t)\right]X_k^{\mathrm{DBO}}',26)
    set_rich_para(p[126],[('The weight decreases from 0.8 to 0.2. With ',False),(r'p=t/T_{\max}',True),(', the implemented RIME component is ',False),(r'X^{\mathrm{best}}+0.56(1-p)\beta\cos(\theta)',True),(', where ',False),(r'\beta\sim\mathcal N(0,1)',True),(' and ',False),(r'\theta\sim\mathcal U(0,2\pi)',True),('. The DBO component is ',False),(r'X+(1-p)kX+\xi\sqrt{(X-X^{\mathrm{worst}})^2}',True),(', where ',False),(r'k\sim\mathcal U(-1,1)',True),(' and ',False),(r'\xi\sim\mathcal U(0,1)',True),('. Equation (26) uses population-individual index k rather than task index i and combines the two candidate components before decoding.',False)],12)
    set_rich_para(p[127],[('Followers use hard-rime puncture with probability ',False),(r'\min\left(1,2e^{-16p^2}\right)',True),('; a Bernoulli mask with probability 0.20 replaces selected task rows by the current best. The alternative foraging update is ',False),(r'X_k^{\mathrm{new}}=X^{\mathrm{best}}+c_1(X_k-L)+c_2(X_k-U)',True),(', with ',False),(r'c_1,c_2\sim\mathcal U(0,1)',True),(', coordinate-wise lower bound ',False),(r'L=(0,0.2)',True),(', and upper bound ',False),(r'U=(2,1)',True),(' broadcast across tasks. This corrects the earlier shorthand that used 0 and 2 for both coordinates. Scouts use ',False),(r'X^{\mathrm{local}}+\tan(\theta)\sqrt{(X-X^{\mathrm{local}})^2}',True),(' with ',False),(r'\theta\sim\mathcal U(-\pi/4,\pi/4)',True),(' when their search fitness is worse than the best, or ',False),(r'X^{\mathrm{best}}+0.035(1-p)\zeta',True),(' otherwise, where ',False),(r'\zeta\sim\mathrm{Cauchy}(0,1)',True),('.',False)],12)

    # Reproducibility and result interpretation.
    set_para(p[135],('The experiments were implemented in Python 3.13.5 and executed as single-process algorithm runs in a Debian GNU/Linux 13 container allocated 56 virtual CPU cores and 4 GB RAM. The principal libraries were NumPy 2.3.5, Pandas 2.2.3, SciPy 1.17.0, Matplotlib 3.10.8, and PyYAML 6.0.3. Scenario seeds start at 20260701 and increase by run. Baseline seeds are derived from the scenario seed and algorithm label. All RDHO component variants use the same RDHO seed label for each scenario, so their initial pseudorandom stream is paired; later draws can diverge when a disabled mechanism changes control flow. The RDHO-full reference in the ablation analysis therefore reproduces the main-experiment solution metrics rather than using an unrelated algorithm-label seed.'),12)
    set_para(p[136],("RIME and DBO provide parent-algorithm reference implementations, not controls that isolate the hybrid operator, because they do not receive RDHO's greedy seed or coordinate refinement. TLBO-HHO [21] and the chaotic-weighted t-distribution sparrow search algorithm (CWTSSA) [35] represent enhanced population-based references, and Greedy Energy-Delay (Greedy-ED) is a lightweight coordinate-wise reference. Every method receives the same task/network scenario and the same metric/reporting implementation."),12)
    set_para(p[137],f'Population size and iteration count are equal for population-based methods, but evaluation budgets are not identical. RDHO-full averages {rdho.nfe_mean:.0f} NFE, compared with 7551 NFE for the population baselines; Greedy-ED uses 361 NFE. Table 5 is an end-to-end solver comparison under configured procedures, not an equal-NFE test of the population operators. RDHO-core provides a closer—but still not perfectly NFE-matched—view of the population stage and retains the greedy seed. Accordingly, no sentence attributes Table 5 significance specifically to the RIME-DBO fusion operator.',12)
    set_para(p[142],('Fig. 2 reports the mean fixed-reference incumbent over 150 population iterations. The curve excludes final coordinate refinement and uses the corrected coordinate-wise bounds in follower foraging. It compares population-search trajectories on a common reporting scale, but initialisation and exact NFE still differ, so it is not interpreted as an isolated test of the RIME-DBO operator.'),12)
    set_para(p[146],f'RDHO-full obtains the lowest mean reporting fitness ({rdho.fitness_mean:.4f} ± {rdho.fitness_std:.4f}); the relative reductions are {imptext}. Its mean QoE, priority-aware fairness, and soft CSR are {rdho.qoe_mean:.4f}, {rdho.fairness_mean:.4f}, and {rdho.csr_mean:.4f}. These results support the complete configured solver only; they do not separate RIME-DBO population updates from greedy seeding and coordinate refinement and do not establish dominance on every raw metric.',12)
    set_para(p[162],('RDHO-full reaches the highest mean aggregate soft CSR. With 40 tasks and three threshold checks per task, its mean CSR of {:.4f} corresponds to approximately {:.1f} satisfied checks out of 120, leaving about {:.1f} binary threshold violations per scenario on average. This aggregate does not reveal whether violations arise from delay, battery-adjusted energy, or AoI, and it weights a marginal exceedance and a severe exceedance equally. It is therefore a comparative soft-QoS diagnostic, not a deployment-level SLA or evidence that every task meets its thresholds. An illegal execution mode or computation-control value is handled separately by decoding and clipping.'.format(rdho.csr_mean,120*rdho.csr_mean,120*(1-rdho.csr_mean))),12)
    set_para(p[168],('The analysis separates the population-search stage from final coordinate refinement. RDHO-core disables only refinement; each remaining variant removes one named mechanism while retaining the same refinement setting. For each scenario, all RDHO variants begin from the same derived RDHO random stream, which improves pairing relative to variant-specific labels. Because altered control flow changes subsequent random-number consumption, the comparison remains an algorithm-configuration analysis rather than a common-random-numbers proof of causal attribution.'),12)

    # Main stats table values and ablation narrative are updated after table construction below.
    set_para(p[185],('Fig. 10. RDHO-full weight sensitivity for QoE, priority-aware per-user utility fairness, and soft CSR.'),10)
    set_para(p[194],('Fig. 12. RDHO-full and baseline min-max normalised multi-metric comparison; energy, delay, and AoI are reversed so that larger radial values are better.'),10)
    set_para(p[196],('Overall, RDHO-full offers the lowest fixed-reference reporting fitness in the implemented end-to-end suite, but raw-metric leaders vary, its evaluation budget is larger, and the RIME-DBO population operator is not isolated from greedy seeding and local refinement. The evidence supports a reproducible weighted trade-off delivered by the complete framework under the stated abstraction, not universal dominance, strict CPU-capacity optimality, or superiority of each internal mechanism.'),12)
    set_para(p[199],f'In 30 paired scenarios, RDHO-full achieves the lowest mean fixed-reference reporting fitness ({rdho.fitness_mean:.4f}) among the implemented complete solvers. Two-sided corrected paired tests, component/configuration analysis, scalability analysis, and sensitivity checks provide complementary evidence. The results concern the complete RDHO-full pipeline; they do not isolate RIME-DBO fusion from greedy initialisation and coordinate refinement. Several baselines remain preferable on individual raw metrics, runtime, or evaluation cost.',12)
    set_para(p[200],('The study remains limited by simulated offline single-epoch task sets, fixed associations and rates, heuristic service-model coefficients without hardware calibration, a non-additive load-adjusted service model rather than capacity-conserving CPU scheduling, an aggregate binary CSR without violation-type or magnitude decomposition, a delay-coupled AoI surrogate, continuous relaxation that imposes an algorithmic neighbourhood on categorical modes, priority-aware rather than priority-neutral fairness, coupled raw-metric and QoE weighting, and a non-NFE-matched end-to-end comparison. Future work should examine calibrated processor sharing, coefficient and SLA sensitivity, decomposed constraint violations, queue-aware AoI, discrete mode operators, priority-neutral fairness, equal-NFE or common-postprocessing baselines, online arrivals and physical testbeds.'),12)

    algorithm_rows=[
      ['Require:','Task/network instance; weights; λ₀, α; population P; iterations T'],
      ['Ensure:','Best solution X* evaluated by reporting fitness and its raw metrics'],
      ['1','Generate half Gaussian and half uniform individuals'],
      ['2','Insert one greedy coordinate seed and up to three perturbations'],
      ['3','Evaluate base metrics and initialise search/reporting incumbents'],
      ['4','for t=1,…,T do'],
      ['5','Recompute parent search fitness with the current λ(t)'],
      ['6','Sort population; assign roles; preserve configured elites'],
      ['7','Generate producer candidates from weighted RIME-/DBO-inspired components'],
      ['8','Generate follower candidates by puncture or bound-aware foraging'],
      ['9','Generate scout candidates by theft or decaying Cauchy mutation'],
      ['10','Round/clip z and clip r to coordinate-specific legal bounds'],
      ['11','Evaluate candidates once to obtain base objective and CSR'],
      ['12','Derive parent/candidate search fitness with the same λ(t)'],
      ['13','Greedily accept candidates with lower current search fitness'],
      ['14','Update the independent reporting-fitness incumbent'],
      ['15','end for'],
      ['16','If enabled, run at most two coordinate-refinement passes'],
      ['17','Return X*, reporting fitness, raw metrics, and NFE'],
    ]
    set_table(doc.tables[1],algorithm_rows,[.75,5.65],8.4)

    # Notation table.
    notation=[
      ['Symbol','Definition'],
      ['𝓜, 𝓖, 𝓒, 𝓝','Device, edge-server, cloud-server, and task sets'],
      ['M, G, C, N','Cardinalities of 𝓜, 𝓖, 𝓒, and 𝓝'],
      ['m(i), g(i), c(i)','Source device and fixed access-edge/cloud mappings for task i'],
      ['zᵢ','Execution mode: 0 local, 1 edge, 2 cloud'],
      ['rᵢ','Normalised computation-control value, 0.2≤rᵢ≤1'],
      ['Lᵢ, Cᵢ','Input bits and required CPU cycles'],
      ['Dᵢ, Eᵢ, Aᵢ','Delay, device-side energy proxy, and single-epoch AoI surrogate'],
      ['Dᵢᵐᵃˣ, Eᵢᵇᵘᵈ, Aᵢᵐᵃˣ','Task-specific delay, nominal energy, and AoI thresholds'],
      ['bᵢ, πᵢ, Δᵢ','Battery ratio, task priority, and update interval'],
      ['fᵢᴸ, fᵢᴱ, fᵢᶜ','Non-additive load-adjusted effective service rates'],
      ['qᵢ, q̄','Priority-weighted task QoE proxy and global task-average QoE'],
      ['q̄ₘ','Per-user mean priority-aware utility'],
      ['J','Jain fairness over active-user mean priority-aware utility'],
      ['CSR','Aggregate soft delay/energy/AoI satisfaction ratio'],
      ['Fbase','Continuous weighted base objective'],
      ['Fsearch','Iteration-dependent internal search fitness'],
      ['Freport','Fixed-reference reporting fitness for one weight vector'],
      ['λ(t), λref','Dynamic search coefficient and fixed reporting coefficient'],
      ['NFE','Number of objective-function evaluations'],
    ]
    set_table(doc.tables[0],notation,[1.65,4.75],8.5)

    # System coefficient table expanded with rationale.
    table2=[
      ['Parameter','Value and interpretation'],
      ['Devices / edge / cloud servers','20 / 4 / 2'],
      ['Task count','40 heterogeneous tasks'],
      ['Device / edge / cloud CPU','0.8-2.2 / 8-18 / 25-40 GHz'],
      ['Device transmit power','0.2-0.8 W'],
      ['Device-edge / edge-cloud rate','8-30 / 60-150 Mbps'],
      ['Device energy coefficient','(0.8-1.4)×10⁻²⁷ J·s²/cycle³ (DVFS proxy)'],
      ['Control affine factors','Local (0.35,0.65), edge (0.40,0.60), cloud (0.45,0.55); dimensionless, monotonic service-control mapping'],
      ['Congestion factors','0.07 / 0.035 / 0.020 per additional local/edge/cloud task; fixed heuristic stress-test factors'],
      ['Fixed edge / cloud delay overhead','0.010 / 0.055 s; simulation-level service overhead'],
      ['Device-active energy overhead','0.015 / 0.010 J/s for edge/cloud service; device-side proxy'],
      ['Calibration status','Inherited from preliminary thesis simulator; fixed before comparison, not fitted to hardware traces or tuned per algorithm'],
    ]
    set_table(doc.tables[2],table2,[2.4,4.0],8.1)

    # Main table.
    main_table=[['Alg.','Reporting fitness','Energy (J)','Delay (s)','AoI surrogate (s)','QoE','Priority-aware fairness','Soft CSR','Time (s)','NFE']]
    for name in ['RDHO','RIME','DBO','TLBO-HHO','CWTSSA','Greedy-ED']:
        r=summary(main_summary,'algorithm',name)
        main_table.append(['RDHO-full' if name=='RDHO' else name,fmt(r.fitness_mean,r.fitness_std),fmt(r.energy_mean,r.energy_std,1),fmt(r.delay_mean,r.delay_std),fmt(r.aoi_mean,r.aoi_std),fmt(r.qoe_mean,r.qoe_std),fmt(r.fairness_mean,r.fairness_std),fmt(r.csr_mean,r.csr_std),fmt(r.runtime_mean,r.runtime_std),f'{r.nfe_mean:.0f}'])
    set_table(doc.tables[5],main_table,[.62,.92,.76,.70,.77,.57,.82,.68,.68,.55],7.1)

    # Stats table, now actually two-sided and includes W/median.
    stats_table=[['Comparison','W','p (2-sided)','Holm p','Median diff.','Signed r_rb','W/T/L']]
    for _,r in stats.iterrows():
        stats_table.append([r.comparison.replace('RDHO vs','RDHO-full vs'),f'{r.w_statistic:.0f}',f'{r.p_value:.2e}',f'{r.p_holm:.2e}',f'{r.median_difference:.4f}',f'{r.rank_biserial:.3f}',f'{int(r.wins)}/{int(r.ties)}/{int(r.losses)}'])
    set_table(doc.tables[6],stats_table,[1.72,.36,.78,.78,.82,.58,.58],7.0)

    # Ablation table: use main RDHO row for full so objective metrics exactly match Table 5.
    ab_table=[['Variant','Reporting fitness','QoE','Priority-aware fairness','Soft CSR','Time (s)','NFE']]
    order=['RDHO-full','RDHO-core','RDHO-w/o dual-source initialization','RDHO-w/o adaptive role allocation','RDHO-w/o elite preservation','RDHO-w/o dynamic penalty']
    labels={'RDHO-w/o dual-source initialization':'w/o dual-source init.','RDHO-w/o adaptive role allocation':'w/o adaptive roles','RDHO-w/o elite preservation':'w/o elite','RDHO-w/o dynamic penalty':'w/o penalty'}
    for name in order:
        r=rdho if name=='RDHO-full' else summary(ab_summary,'algorithm',name)
        ab_table.append([labels.get(name,name),fmt(r.fitness_mean,r.fitness_std),fmt(r.qoe_mean,r.qoe_std),fmt(r.fairness_mean,r.fairness_std),fmt(r.csr_mean,r.csr_std),fmt(r.runtime_mean,r.runtime_std),f'{r.nfe_mean:.0f}'])
    set_table(doc.tables[7],ab_table,[1.38,.98,.68,.88,.72,.72,.55],7.7)

    # Update ablation narrative from paired stats.
    core=summary(ab_summary,'algorithm','RDHO-core')
    dual=summary(ab_summary,'algorithm','RDHO-w/o dual-source initialization')
    def ast(name): return abs_stats[abs_stats.variant==name].iloc[0]
    corest=ast('RDHO-core'); dualst=ast('RDHO-w/o dual-source initialization')
    nonsig=[]
    for n,label in [('RDHO-w/o adaptive role allocation','adaptive role allocation'),('RDHO-w/o elite preservation','elite preservation'),('RDHO-w/o dynamic penalty','dynamic penalty')]:
        rr=ast(n); nonsig.append((label,rr.p_holm))
    adapt=ast('RDHO-w/o adaptive role allocation'); elite=ast('RDHO-w/o elite preservation'); pen=ast('RDHO-w/o dynamic penalty')
    set_para(p[172],f'Coordinate refinement reduces mean reporting fitness from {core.fitness_mean:.4f} for RDHO-core to {rdho.fitness_mean:.4f} for RDHO-full (median paired difference {corest["median"]:.4f}; Holm-adjusted two-sided p={corest["p_holm"]:.2e}). Removing dual-source initialisation changes the mean to {dual.fitness_mean:.4f} (adjusted p={dualst["p_holm"]:.2e}). The adjusted p-values for removing adaptive role allocation, elite preservation, and dynamic penalty are {adapt["p_holm"]:.3f}, {elite["p_holm"]:.3f}, and {pen["p_holm"]:.3f}, respectively, so no independent necessity is claimed for those controls. The complete output is available at results/summary/ablation_wilcoxon_results.csv.',12)

    # Scalability / sensitivity tables and narrative.
    scale_table=[['Tasks','Reporting fitness','Soft CSR','Time (s)','NFE']]
    for _,r in scale.sort_values('task_number').iterrows(): scale_table.append([f'{int(r.task_number)}',fmt(r.fitness_mean,r.fitness_std),fmt(r.csr_mean,r.csr_std),fmt(r.runtime_mean,r.runtime_std),f'{r.nfe_mean:.0f}'])
    set_table(doc.tables[8],scale_table,[.75,1.35,1.0,1.0,.8],8.5)
    first=scale.sort_values('task_number').iloc[0]; last=scale.sort_values('task_number').iloc[-1]
    set_para(p[178],f'Mean reporting fitness changes from {first.fitness_mean:.4f} at {int(first.task_number)} tasks to {last.fitness_mean:.4f} at {int(last.task_number)} tasks, while soft CSR changes from {first.csr_mean:.4f} to {last.csr_mean:.4f}. Runtime grows from {first.runtime_mean:.3f} s to {last.runtime_mean:.3f} s. These results use the corrected bound-aware foraging update and describe empirical scaling for the tested range rather than an asymptotic complexity proof. The 40-task scalability row uses the ten-seed scalability subset and therefore does not numerically reproduce the 30-run main-experiment mean.',12)
    weight_table=[['Set','Weights (E,D,A,Q,J)','Setting-specific fitness','QoE','Priority-aware fairness','Soft CSR','Time (s)']]
    for setting in ['S1','S2','S3','S4','S5']:
        r=summary(weight,'setting',setting); weight_table.append([setting,str(r.weights),fmt(r.fitness_mean,r.fitness_std),fmt(r.qoe_mean,r.qoe_std),fmt(r.fairness_mean,r.fairness_std),fmt(r.csr_mean,r.csr_std),fmt(r.runtime_mean,r.runtime_std)])
    set_table(doc.tables[9],weight_table,[.42,1.45,1.02,.66,.88,.70,.70],7.4)
    set_para(p[181],f'Across S1-S5, mean QoE ranges from {weight.qoe_mean.min():.4f} to {weight.qoe_mean.max():.4f}, priority-aware fairness from {weight.fairness_mean.min():.4f} to {weight.fairness_mean.max():.4f}, and soft CSR from {weight.csr_mean.min():.4f} to {weight.csr_mean.max():.4f}. The relatively narrow raw-metric ranges are more informative than differences in setting-specific fitness, which partly reflect direct reweighting.',12)
    penalty_table=[['λ₀','α','Reporting fitness','QoE','Priority-aware fairness','Soft CSR','Time (s)']]
    for _,r in penalty.sort_values(['lambda0','alpha']).iterrows(): penalty_table.append([f'{r.lambda0:.1f}',f'{r.alpha:.1f}',fmt(r.fitness_mean,r.fitness_std),fmt(r.qoe_mean,r.qoe_std),fmt(r.fairness_mean,r.fairness_std),fmt(r.csr_mean,r.csr_std),fmt(r.runtime_mean,r.runtime_std)])
    set_table(doc.tables[10],penalty_table,[.48,.48,1.08,.68,.9,.72,.72],7.6)
    set_para(p[186],f'Across the nine λ₀-α combinations, mean soft CSR ranges from {penalty.csr_mean.min():.4f} to {penalty.csr_mean.max():.4f}, and fixed-reference reporting fitness ranges from {penalty.fitness_mean.min():.4f} to {penalty.fitness_mean.max():.4f}. The default weight vector and λref=1 are fixed, so these values are directly comparable and reflect solutions returned under different dynamic search schedules.',12)

    # Replace regenerated figures once present.
    figs={143:(REPO/'results/figures/convergence_curve.png',4.85),148:(REPO/'results/figures/energy_comparison.png',4.85),151:(REPO/'results/figures/delay_comparison.png',4.85),154:(REPO/'results/figures/aoi_comparison.png',4.85),157:(REPO/'results/figures/qoe_fairness_comparison.png',4.85),160:(REPO/'results/figures/csr_comparison.png',4.85),170:(REPO/'results/figures/ablation_study_multicolor.png',4.85),176:(REPO/'results/figures/scalability.png',4.85),184:(REPO/'results/sensitivity/figures/weight_sensitivity_qoe_fairness_csr.png',4.85),189:(REPO/'results/sensitivity/figures/penalty_sensitivity_heatmaps.png',4.85),193:(REPO/'results/figures/radar_chart.png',3.82)}
    for idx,(path,w) in figs.items():
        if path.exists(): replace_picture(p[idx],path,w)

    # Global terminology cleanup without destroying drawing-bearing runs.
    replacements={'resource-ratio':'computation-control','resource ratio':'computation-control value','resource ratios':'computation-control values','User fairness':'Priority-aware fairness','user fairness':'priority-aware fairness'}
    replace_terms_preserving_drawings(doc, replacements)

    # References: compact and remove pagination controls to avoid a nearly blank final page.
    for para in doc.paragraphs[212:]:
        para.paragraph_format.space_before=Pt(0); para.paragraph_format.space_after=Pt(0); para.paragraph_format.line_spacing=1.0
        para.paragraph_format.keep_with_next=False; para.paragraph_format.keep_together=False; para.paragraph_format.page_break_before=False
        for run in para.runs:
            run.font.size=Pt(9.3); run.font.name='Times New Roman'; run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'),'Times New Roman')
    # Remove unnecessary empty body paragraphs while preserving fixed paragraph count (use zero spacing).
    for idx in [95,143,148,151,154,157,160,170,176,184,188,189,193,211]:
        p[idx].paragraph_format.space_before=Pt(0); p[idx].paragraph_format.space_after=Pt(0)

    doc.core_properties.subject='Second major revision: corrected bound-aware foraging, unified symbols/seeds, conservative attribution'
    doc.save(INTERMEDIATE)
    patch_comments(INTERMEDIATE,OUTPUT)
    print(OUTPUT)

if __name__=='__main__': main()
