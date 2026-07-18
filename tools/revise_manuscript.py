from __future__ import annotations

import copy
import json
import math
import shutil
import subprocess
import tempfile
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from scipy.stats import friedmanchisquare


REPO = Path(__file__).resolve().parents[1]
SOURCE_DOCX = Path('/mnt/data/0712-review_content_revised (1)(3).docx')
OUTPUT_DOCX = Path('/mnt/data/RDHO_SCI_revision_marked.docx')

W_NS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
M_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'


def _comment_ids(paragraph) -> list[str]:
    ids: list[str] = []
    for element in paragraph._p.xpath('.//*[local-name()="commentRangeStart" or local-name()="commentReference"]'):
        value = element.get(qn('w:id'))
        if value is not None and value not in ids:
            ids.append(value)
    return ids


def _clear_preserve_ppr(paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn('w:pPr'):
            paragraph._p.remove(child)


def _add_comment_start(paragraph, comment_id: str) -> None:
    start = OxmlElement('w:commentRangeStart')
    start.set(qn('w:id'), comment_id)
    paragraph._p.append(start)


def _add_comment_end_reference(paragraph, comment_id: str) -> None:
    end = OxmlElement('w:commentRangeEnd')
    end.set(qn('w:id'), comment_id)
    paragraph._p.append(end)
    run = OxmlElement('w:r')
    rpr = OxmlElement('w:rPr')
    style = OxmlElement('w:rStyle')
    style.set(qn('w:val'), 'CommentReference')
    rpr.append(style)
    run.append(rpr)
    reference = OxmlElement('w:commentReference')
    reference.set(qn('w:id'), comment_id)
    run.append(reference)
    paragraph._p.append(run)


def _format_run(run, size: float | None = 12, bold: bool | None = None, italic: bool | None = None, highlight: bool = True, font: str = 'Times New Roman') -> None:
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def set_paragraph(paragraph, text: str, *, size: float | None = 12, bold: bool | None = None, italic: bool | None = None, highlight: bool = True, alignment=None) -> None:
    ids = _comment_ids(paragraph)
    _clear_preserve_ppr(paragraph)
    for cid in ids:
        _add_comment_start(paragraph, cid)
    if text:
        run = paragraph.add_run(text)
        _format_run(run, size=size, bold=bold, italic=italic, highlight=highlight)
    for cid in reversed(ids):
        _add_comment_end_reference(paragraph, cid)
    if alignment is not None:
        paragraph.alignment = alignment


def set_math_text_paragraph(paragraph, text: str, *, size: float = 12, highlight: bool = True) -> None:
    set_paragraph(paragraph, text, size=size, highlight=highlight, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for run in paragraph.runs:
        run.font.name = 'Cambria Math'
        run._element.get_or_add_rPr().rFonts.set(qn('w:eastAsia'), 'Cambria Math')


def _highlight_math_node(node) -> None:
    for math_run in node.xpath('.//*[local-name()="r" and namespace-uri()="%s"]' % M_NS):
        rpr = next((c for c in math_run if c.tag == qn('m:rPr')), None)
        if rpr is None:
            rpr = OxmlElement('m:rPr')
            math_run.insert(0, rpr)
        wrpr = next((c for c in rpr if c.tag == qn('w:rPr')), None)
        if wrpr is None:
            wrpr = OxmlElement('w:rPr')
            rpr.append(wrpr)
        highlight = OxmlElement('w:highlight')
        highlight.set(qn('w:val'), 'yellow')
        wrpr.append(highlight)


def build_math_nodes(formulas: list[str]) -> list:
    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)
        md = tmp / 'equations.md'
        docx = tmp / 'equations.docx'
        md.write_text('\n\n'.join(f'$$\n{formula}\n$$' for formula in formulas), encoding='utf-8')
        subprocess.run(['pandoc', str(md), '-o', str(docx)], check=True)
        generated = Document(docx)
        nodes = []
        for paragraph in generated.paragraphs:
            found = paragraph._p.xpath('.//*[local-name()="oMath"]')
            if found:
                node = copy.deepcopy(found[0])
                _highlight_math_node(node)
                nodes.append(node)
        if len(nodes) != len(formulas):
            raise RuntimeError(f'Expected {len(formulas)} equations, found {len(nodes)}')
        return nodes


def _set_equation_tab_stops(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    for existing in list(ppr.findall(qn('w:tabs'))):
        ppr.remove(existing)

    tabs = OxmlElement('w:tabs')
    center_tab = OxmlElement('w:tab')
    center_tab.set(qn('w:val'), 'center')
    center_tab.set(qn('w:pos'), '4860')
    tabs.append(center_tab)

    right_tab = OxmlElement('w:tab')
    right_tab.set(qn('w:val'), 'right')
    right_tab.set(qn('w:pos'), '9720')
    tabs.append(right_tab)

    pstyle = ppr.find(qn('w:pStyle'))
    if pstyle is None:
        ppr.insert(0, tabs)
    else:
        ppr.insert(list(ppr).index(pstyle) + 1, tabs)


def set_equation(paragraph, math_node, number: int) -> None:
    ids = _comment_ids(paragraph)
    _clear_preserve_ppr(paragraph)
    _set_equation_tab_stops(paragraph)
    for cid in ids:
        _add_comment_start(paragraph, cid)

    tab1 = OxmlElement('w:r')
    tab1.append(OxmlElement('w:tab'))
    paragraph._p.append(tab1)
    paragraph._p.append(copy.deepcopy(math_node))
    paragraph.add_run('\t')
    number_run = paragraph.add_run(f'({number})')
    _format_run(number_run, size=11, highlight=True)

    for cid in reversed(ids):
        _add_comment_end_reference(paragraph, cid)


def set_display_equation(paragraph, math_node) -> None:
    ids = _comment_ids(paragraph)
    _clear_preserve_ppr(paragraph)
    for cid in ids:
        _add_comment_start(paragraph, cid)
    paragraph._p.append(copy.deepcopy(math_node))
    for cid in reversed(ids):
        _add_comment_end_reference(paragraph, cid)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def replace_picture(paragraph, image_path: Path, width_inches: float) -> None:
    _clear_preserve_ppr(paragraph)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))


def _set_cell_border(cell, **edges) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in('w:tcBorders')
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge not in edges:
            continue
        edge_data = edges[edge]
        tag = 'w:' + edge
        element = tcBorders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tcBorders.append(element)
        for key, value in edge_data.items():
            element.set(qn('w:' + key), str(value))


def ensure_table_size(table, rows: int, cols: int) -> None:
    while len(table.rows) < rows:
        table.add_row()
    while len(table.columns) < cols:
        table.add_column(Inches(0.6))


def set_cell(cell, text: str, *, header: bool = False, size: float = 9, highlight: bool = True) -> None:
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(str(text))
    _format_run(run, size=size, bold=header, highlight=highlight)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_data(table, data: list[list[str]], widths: list[float] | None = None, font_size: float = 9) -> None:
    ensure_table_size(table, len(data), max(len(row) for row in data))
    table.autofit = False
    if widths:
        for idx, width in enumerate(widths):
            for cell in table.columns[idx].cells:
                cell.width = Inches(width)
    for r, row in enumerate(data):
        for c in range(len(table.columns)):
            value = row[c] if c < len(row) else ''
            set_cell(table.cell(r, c), value, header=(r == 0), size=font_size)
    # Blank any unused rows while retaining document structure.
    for r in range(len(data), len(table.rows)):
        for c in range(len(table.columns)):
            set_cell(table.cell(r, c), '', size=font_size, highlight=False)

    nil = {'val': 'nil'}
    single = {'val': 'single', 'sz': '8', 'space': '0', 'color': '000000'}
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell, top=nil, bottom=nil, left=nil, right=nil, insideH=nil, insideV=nil)
    for cell in table.rows[0].cells:
        _set_cell_border(cell, top=single, bottom=single)
    for cell in table.rows[len(data) - 1].cells:
        _set_cell_border(cell, bottom=single)


def fmt(mean: float, std: float, decimals: int) -> str:
    return f'{mean:.{decimals}f} ± {std:.{decimals}f}'


def ablation_label(name: str) -> str:
    labels = {
        'RDHO-full': 'RDHO-full',
        'RDHO-core': 'RDHO-core',
        'RDHO-w/o dual-source initialization': 'w/o dual-source init.',
        'RDHO-w/o adaptive role allocation': 'w/o adaptive roles',
        'RDHO-w/o elite preservation': 'w/o elite',
        'RDHO-w/o dynamic penalty': 'w/o penalty',
    }
    return labels.get(name, name.replace('RDHO-', ''))


def summary_row(df: pd.DataFrame, key_col: str, key: str) -> pd.Series:
    row = df[df[key_col] == key]
    if row.empty:
        raise KeyError(key)
    return row.iloc[0]


def best_name(df: pd.DataFrame, metric: str, higher: bool) -> str:
    grouped = df.groupby('algorithm')[metric].mean()
    return str(grouped.idxmax() if higher else grouped.idxmin())


def load_results():
    main_raw = pd.read_csv(REPO / 'results/raw/main_30_raw_results.csv')
    main_summary = pd.read_csv(REPO / 'results/summary/main_30_summary_mean_std.csv')
    stats = pd.read_csv(REPO / 'results/summary/wilcoxon_fitness_results.csv')
    ablation = pd.read_csv(REPO / 'results/summary/ablation_30_summary_mean_std.csv')
    scale = pd.read_csv(REPO / 'results/summary/scalability_summary_mean_std.csv')
    weight = pd.read_csv(REPO / 'results/sensitivity/summary/weight_sensitivity_summary_mean_std.csv')
    penalty = pd.read_csv(REPO / 'results/sensitivity/summary/dynamic_penalty_sensitivity_summary_mean_std.csv')
    return main_raw, main_summary, stats, ablation, scale, weight, penalty


def main() -> None:
    main_raw, main_summary, stats, ablation, scale, weight, penalty = load_results()
    doc = Document(SOURCE_DOCX)
    p = doc.paragraphs

    rdho = summary_row(main_summary, 'algorithm', 'RDHO')
    baseline_names = [name for name in ['RIME', 'DBO', 'TLBO-HHO', 'CWTSSA', 'Greedy-ED'] if name in set(main_summary['algorithm'])]
    improvements = {
        name: 100.0 * (summary_row(main_summary, 'algorithm', name)['fitness_mean'] - rdho['fitness_mean']) / summary_row(main_summary, 'algorithm', name)['fitness_mean']
        for name in baseline_names
    }
    improvement_text = ', '.join(f'{improvements[name]:.1f}% versus {name}' for name in baseline_names)

    # Title, abstract, introduction, and related work.
    set_paragraph(p[0], 'RIME-DBO-Based QoE- and Fairness-Aware Task Offloading in Mobile Edge Computing', size=16, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    set_paragraph(p[6], (
        'Mobile edge computing (MEC) reduces service distance for resource-constrained devices, but heterogeneous workloads require decisions that jointly consider efficiency, information freshness, perceived service quality, and equitable user outcomes. '
        'This paper proposes RIME-DBO Optimisation (RDHO), a mixed discrete-continuous scheduler that selects local, edge, or cloud execution together with a bounded normalised computation-control value. '
        'The model combines mobile-device energy, processing delay, average Age of Information (AoI), a priority-weighted Quality of Experience (QoE) proxy, and user-level Jain fairness. '
        'An iteration-dependent penalty guides the search, while a fixed-reference reporting fitness ensures comparable final evaluation across algorithms and sensitivity settings. '
        'Thirty paired simulations, statistical tests, component ablations, scalability analysis, and parameter-sensitivity experiments show that RDHO provides the strongest overall reporting objective in the implemented comparison suite, although individual baselines can remain preferable on specific raw metrics or computational cost.'
    ), size=12)
    set_paragraph(p[7], 'Keywords: Mobile Edge Computing; task offloading; Quality of Experience; priority-aware fairness; metaheuristic optimisation.', size=12)
    set_paragraph(p[9], 'Latency-sensitive sensor, Internet of Things (IoT), and 5G/6G services increasingly generate heterogeneous workloads close to end users. Mobile Edge Computing (MEC) shortens the service path by complementing resource-limited devices with nearby edge and remote cloud computing resources [1,2].', size=12)
    set_paragraph(p[10], 'The scheduling problem considered here is to choose local, edge, or cloud execution for every task and to determine a bounded normalised computation-control value for the selected mode. The decisions are coupled through shared-node load, transmission delay, device-side energy, and service-quality thresholds, producing a nonlinear mixed discrete-continuous search space.', size=12)
    set_paragraph(p[11], 'Most MEC offloading studies emphasise energy and delay, while freshness-aware studies additionally consider AoI [3,4]. These indicators remain necessary but do not fully describe how acceptable the service is to heterogeneous users or whether users receive comparably satisfactory outcomes [5-7].', size=12)
    set_paragraph(p[12], 'This work therefore integrates a model-based QoE proxy and user-level QoE fairness with energy, delay, and average AoI. RDHO is positioned as a user-centric offloading scheduler for this five-metric objective rather than as a universally superior optimiser for all MEC formulations.', size=12)
    set_paragraph(p[13], 'The main contributions are fourfold. First, we formulate a cloud-edge-device offloading model whose decision vector contains an execution mode and normalised computation-control value for each task, with fixed access-edge and cloud associations. Second, we develop RDHO by combining dual-source initialisation, adaptive producer/follower/scout roles, elite preservation, RIME-DBO updates, and an optional coordinate-wise local refinement. Third, we distinguish the iteration-dependent search fitness from a fixed-reference reporting fitness and compare parents and candidates under the same current penalty coefficient, removing an otherwise ambiguous evaluation inconsistency. Fourth, we conduct paired 30-run comparisons, corrected non-parametric tests, isolated ablations, scalability and sensitivity studies, and release the complete reproducibility artefact.', size=12)
    set_paragraph(p[14], f'Across the paired benchmark, RDHO obtains a mean fixed-reference reporting fitness of {rdho["fitness_mean"]:.4f} and improves that objective by {improvement_text}. Its advantage is interpreted as an overall weighted trade-off; the raw energy, delay, AoI, QoE, fairness, runtime, and number of function evaluations are reported separately so that metric-specific trade-offs remain visible.', size=12)
    set_paragraph(p[15], 'The remainder of this paper is organised as follows. Section 2 reviews related work. Section 3 defines the system model and objectives. Section 4 presents RDHO and its evaluation semantics. Section 5 reports the experiments, and Section 6 concludes the paper.', size=12)
    set_paragraph(p[17], 'Following recent MEC surveys, related work is organised into energy-delay optimisation, learning-based scheduling, freshness-aware scheduling, QoE and fairness modelling, and metaheuristic optimisation [8-10]. This classification separates the technical foundations of offloading from the user-centric and algorithmic gap addressed here.', size=12)
    set_paragraph(p[18], 'Energy- and delay-aware offloading established joint management of execution placement and computational resources as a central MEC problem [1-3,11]. Later studies extended this perspective to collaborative edge-cloud systems, digital-twin service placement, and security-aware vehicular environments [12-14]. The present implementation adopts fixed communication rates and associations and optimises the execution mode and a bounded normalised computation-control value rather than radio bandwidth or server identity.', size=12)
    set_paragraph(p[19], 'Learning-based offloading adapts to changing channel and resource states. Deep and federated reinforcement-learning approaches have been applied to online computation offloading, collaborative vehicular computing, and joint resource management [15-17]. Their online adaptability is attractive, but training cost and data dependence motivate complementary training-free optimisers for transparent offline or epoch-level decisions.', size=12)
    set_paragraph(p[20], 'Freshness-aware MEC optimisation uses AoI to complement latency [18-20]. The earlier TLBO-HHO study combined energy, delay, and freshness in a hybrid metaheuristic framework [21]. In contrast, the present work additionally includes a priority-weighted QoE proxy and fairness over per-user mean QoE.', size=12)
    set_paragraph(p[21], 'QoE research links measurable service performance to perceived utility, while Jain\'s index and proportional fairness are established tools for allocation equity [5,6,22-25]. Fairness-aware offloading has also been examined in cooperative and vehicular systems [7,26]. Here, task QoE values are first aggregated by source device, and fairness is computed across active users rather than across individual tasks.', size=12)
    set_paragraph(p[22], 'Because the decision space combines categorical execution modes, continuous computation-control values, load-dependent performance, and nonlinear objective terms, population-based metaheuristics remain practical for medium-sized simulated instances. RIME and DBO provide complementary exploration and exploitation mechanisms [33,34], but hybridisation must be supported by controlled component studies rather than inferred solely from a final score.', size=12)
    set_paragraph(p[23], 'The resulting gap is a reproducible scheduler that jointly reports technical efficiency, freshness, model-based QoE, and user-level fairness while clearly separating search guidance from final cross-algorithm evaluation. RDHO addresses this gap under the implemented fixed-association, fixed-rate, simulated MEC setting; the claims are limited to the stated comparison suite and objective.', size=12)

    # System model and objective.
    set_paragraph(p[26], 'The considered system uses a three-tier cloud-edge-device architecture. Let M, G, C, and N denote the sets of mobile devices, access-edge servers, cloud servers, and tasks, respectively. Each device has a fixed access edge during one scheduling epoch, and each edge is mapped to one cloud server. The original architecture in Fig. 1 illustrates these local, edge, and cloud execution paths.', size=12)
    set_paragraph(p[27], 'Task i is generated by source device m(i). If edge execution is chosen, it is processed at the fixed access edge g(i); if cloud execution is chosen, it is forwarded through g(i) to the mapped cloud c(i). Server selection and communication-rate allocation are therefore outside the optimisation scope.', size=12)
    set_paragraph(p[29], 'Fig. 1. Three-tier cloud-edge-device MEC architecture and fixed forwarding paths.', size=10)
    set_paragraph(p[30], 'Task i is represented by the tuple', size=12)
    set_math_text_paragraph(p[31], 'τᵢ = (Lᵢ, Cᵢ, Dᵢᵐᵃˣ, Aᵢᵐᵃˣ, Eᵢᵇᵘᵈ, Δᵢ, bᵢ, πᵢ)', size=12)
    set_paragraph(p[32], 'where Lᵢ is the input size, Cᵢ is the required CPU-cycle count, Dᵢᵐᵃˣ and Aᵢᵐᵃˣ are the delay and AoI thresholds, Eᵢᵇᵘᵈ is the nominal energy budget, Δᵢ is the update interval, bᵢ is the residual-battery ratio, and πᵢ is the task priority.', size=12)
    set_paragraph(p[33], 'Each solution contains two variables per task: zᵢ selects local, edge, or cloud execution, and rᵢ is a normalised computation-control value. Tasks are independent, result-download cost is neglected, and rates and associations remain constant within an epoch. Shared-node contention is represented through a load-dependent effective-frequency abstraction.', size=12)
    set_paragraph(p[35], 'For local execution, the effective frequency assigned to task i is', size=12)
    set_paragraph(p[37], 'For edge execution, the corresponding effective frequency is', size=12)
    set_paragraph(p[39], 'For cloud execution, the effective frequency is', size=12)
    set_paragraph(p[41], 'Using fixed device-edge and edge-cloud rates, the end-to-end processing delay is', size=12)
    set_paragraph(p[43], 'The load terms in Eqs. (1)-(3) are the numbers of tasks using the corresponding device, edge, and cloud nodes. The implemented load-sharing coefficients are fixed, while δᴱ=0.010 s and δᶜ=0.055 s are service overheads. This abstraction bounds each task\'s effective frequency but is not an explicit sum-constrained CPU-frequency allocator.', size=12)
    set_paragraph(p[44], 'Local device energy follows the dynamic-voltage-and-frequency-scaling model', size=12)
    set_paragraph(p[46], 'For edge offloading, device-side energy includes uplink transmission and an active-service overhead proportional to edge execution time:', size=12)
    set_paragraph(p[48], 'For cloud offloading, device-side energy includes uplink transmission and an active-service overhead over backhaul and cloud execution time:', size=12)
    set_paragraph(p[50], 'The coefficients 0.015 J/s and 0.010 J/s are simulation-level device-active overhead factors. Powered edge/cloud infrastructure energy is not modelled; the reported energy metric is therefore a device-side proxy and should not be interpreted as total system energy.', size=12)
    set_paragraph(p[51], 'For periodic updates, the implemented average-AoI approximation is half an update interval plus the service delay:', size=12)
    set_paragraph(p[53], 'Equation (8) represents average rather than peak AoI and is applied identically to every algorithm and sensitivity setting.', size=12)

    set_paragraph(p[56], 'The study uses a transparent QoE proxy rather than claiming a user-validated subjective score. It combines delay, energy, and freshness satisfaction through exponential utilities and scales the result by task priority.', size=12)
    set_paragraph(p[57], '(1) Delay satisfaction.', size=12)
    set_paragraph(p[58], 'Delay satisfaction decays relative to the task-specific delay threshold:', size=12)
    set_paragraph(p[60], 'A delay equal to its threshold produces e⁻¹ satisfaction; lower delays receive higher scores and larger violations decay smoothly rather than causing a discontinuous objective.', size=12)
    set_paragraph(p[61], '(2) Energy satisfaction.', size=12)
    set_paragraph(p[62], 'Energy satisfaction is normalised by the task\'s nominal energy budget:', size=12)
    set_paragraph(p[64], 'Residual battery is used in the soft energy-satisfaction check in Section 3.5, while the continuous QoE energy utility uses the nominal budget so that the utility and threshold check remain distinguishable.', size=12)
    set_paragraph(p[65], '(3) Freshness satisfaction.', size=12)
    set_paragraph(p[66], 'AoI satisfaction is defined analogously using the task-specific AoI threshold:', size=12)
    set_paragraph(p[68], '(4) Comprehensive QoE proxy.', size=12)
    set_paragraph(p[69], 'The task-level QoE proxy is the priority-weighted combination', size=12)
    set_paragraph(p[71], 'The coefficients 0.45, 0.30, and 0.25 are fixed for all algorithms. Priority πᵢ lies in [0,1], and the final score is clipped to [0,1]. These coefficients define the implemented utility and are not presented as universally calibrated human-perception parameters.', size=12)
    set_paragraph(p[73], '3.4 User-Level Fairness Measurement', size=12, bold=True)
    set_paragraph(p[74], 'For each active user, task-level QoE values are first averaged over that user\'s generated tasks. Aggregating before computing fairness prevents users with more generated tasks from receiving disproportionate weight.', size=12)
    set_paragraph(p[75], 'Jain\'s Fairness Index is then evaluated over the active-user QoE means [23]:', size=12)
    set_paragraph(p[77], 'Only users with at least one task in the generated epoch are included. J=1 means equal per-user mean QoE, while a lower value indicates a less even distribution of model-based service satisfaction.', size=12)

    set_paragraph(p[79], 'The optimisation combines energy, delay, average AoI, QoE, and user-level fairness. Technical metrics are normalised by task-specific budgets or thresholds before aggregation.', size=12)
    set_paragraph(p[80], '(1) Normalised components and base objective.', size=12)
    set_paragraph(p[81], 'The normalised energy, delay, and AoI components are', size=12)
    set_paragraph(p[83], 'In the order energy, delay, AoI, QoE, and fairness, the default objective weights are (0.15, 0.15, 0.20, 0.25, 0.25) and sum to one.', size=12)
    set_paragraph(p[84], 'A single scalar objective is used because the scheduler must return one deployable decision for each epoch rather than a Pareto set. The base objective is', size=12)
    set_paragraph(p[86], 'Lower values are better. QoE and fairness enter as deficits. Soft service satisfaction is reported separately from the continuous objective.', size=12)
    set_paragraph(p[87], 'The soft constraint-satisfaction ratio is', size=12)
    set_paragraph(p[89], 'where the energy check uses the battery-adjusted task budget. The execution-mode domain is', size=12)
    set_paragraph(p[91], 'and the normalised computation-control value is bounded by', size=12)
    set_paragraph(p[93], 'The aggregate soft-violation rate is', size=12)
    set_paragraph(p[95], '', size=12)
    set_paragraph(p[96], 'RDHO uses an iteration-dependent search fitness', size=12)
    set_paragraph(p[98], 'with the dynamic coefficient', size=12)
    set_paragraph(p[100], 'At every iteration, both the current parent population and new candidates are scored with the same λ(t) before greedy selection. This avoids comparing objective values produced under different penalty coefficients.', size=12)
    set_paragraph(p[101], 'For final tables and cross-setting comparisons, every returned solution is re-evaluated with the common fixed-reference reporting fitness', size=12)
    set_paragraph(p[103], 'where the fixed reporting coefficient is', size=12)
    set_paragraph(p[105], '', size=12)
    set_paragraph(p[106], 'Consequently, the tabled reporting fitness is not the final-iteration dynamic search fitness. The former is comparable across algorithms and penalty settings; the latter is an internal search-guidance value. Soft CSR below one indicates threshold violations, not an invalid execution-mode or computation-control encoding.', size=12)
    set_paragraph(p[107], '3.6 Computational Characteristics', size=12, bold=True)
    set_paragraph(p[108], 'The model is a nonconvex mixed discrete-continuous optimisation problem. Even before optimising computation-control values, N tasks produce 3ᴺ possible execution-mode assignments.', size=12)
    set_paragraph(p[109], 'The objective additionally couples tasks through node loads and contains exponential utilities, user-level aggregation, indicator-based CSR, and a dynamic penalty. Exhaustive enumeration therefore grows exponentially and derivative-based methods are not directly applicable to the complete formulation.', size=12)
    set_paragraph(p[110], 'These characteristics motivate a population-based search that can explore categorical modes, refine bounded continuous ratios, and retain a directly deployable incumbent without claiming a formal complexity reduction not established by the implemented model.', size=12)

    # Algorithm.
    set_paragraph(p[112], 'RDHO addresses the five-metric problem in Section 3 by combining RIME-style exploration with DBO-style role-based exploitation. The optimiser maintains both an iteration-dependent search ranking and a fixed-reference reporting incumbent.', size=12)
    set_paragraph(p[113], 'RIME supplies broad perturbation and best-solution puncture mechanisms [33], which explore alternative local/edge/cloud modes and reuse promising task-level decisions.', size=12)
    set_paragraph(p[114], 'DBO supplies rolling, foraging, and theft-inspired movements [34], which refine computation-control values and help escape poor regions. All candidate variables are clipped to the legal mode and ratio bounds before evaluation.', size=12)
    set_paragraph(p[115], 'The fusion is role-conditioned: producers emphasise broad hybrid exploration, followers alternate between best-guided puncture and foraging, and scouts apply theft or Cauchy perturbations. Roles are recomputed from current search fitness and population diversity.', size=12)
    set_paragraph(p[117], 'Each individual is an N×2 matrix containing the pair (zᵢ,rᵢ) for every task. The mode component is rounded and clipped to {0,1,2}; the resource component is clipped to [0.2,1.0]. Fixed mappings determine the edge and cloud used by an offloaded task.', size=12)
    set_paragraph(p[118], 'No server-index variable or post-hoc capacity-repair operator is used. Load-dependent effective frequencies in Eqs. (1)-(3) provide the shared-resource abstraction, so the manuscript and code use the same two-variable encoding.', size=12)
    set_paragraph(p[119], 'A single evaluation returns the base objective, soft CSR, fixed-reference reporting fitness, and raw metrics. Search fitness at any iteration is then derived from the same base objective and CSR, avoiding unnecessary re-evaluation and keeping search and reporting histories distinct.', size=12)
    set_paragraph(p[121], 'The initial population contains equal Gaussian and uniform subsets. The first individual is a greedy seed produced by coordinate-wise mode/resource trials, and up to three nearby perturbations are inserted to improve early solution quality while retaining global diversity.', size=12)
    set_paragraph(p[122], 'At iteration t, adaptive producer and scout ratios depend on progress and population diversity; followers receive the remaining population subject to a minimum proportion. The top 10% of individuals are protected as elites.', size=12)
    set_paragraph(p[123], 'For producers, the RIME and DBO components are combined through the adaptive fusion weight', size=12)
    set_paragraph(p[126], 'The weight decreases from 0.8 to 0.2, shifting emphasis from RIME exploration toward DBO refinement. The update acts jointly on execution-mode and computation-control coordinates and is clipped before evaluation.', size=12)
    set_paragraph(p[127], 'Followers use hard-rime puncture early with a decaying probability and otherwise apply a DBO-style best-guided foraging update. Scouts use tangent-based theft when their search fitness is worse than the current best, or a decaying Cauchy perturbation otherwise.', size=12)
    set_paragraph(p[128], 'Greedy replacement compares every candidate with its parent under the same λ(t). The fixed-reference incumbent is updated independently. RDHO-full then performs at most two coordinate-refinement passes over three execution modes and five computation-control candidates per task; RDHO-core omits this final stage, and NFE is reported to make its extra cost explicit.', size=12)
    set_paragraph(p[130], 'Algorithm 1 summarises the implemented RDHO procedure.', size=12)
    set_paragraph(p[131], 'Algorithm 1. RDHO for MEC task offloading and computation-control optimisation', size=10)
    set_paragraph(p[132], 'The pseudocode explicitly separates current-penalty selection, fixed-reference incumbent tracking, and optional local refinement. The returned solution is evaluated with the fixed-reference reporting objective, while the dynamic search objective is retained only as an internal diagnostic.', size=12)

    # Experimental setup.
    set_paragraph(p[135], 'The experiments were implemented in Python 3.13 and executed as single-process algorithm runs in a Linux container. The reproducibility artefact contains source code, configurations, validation tests, raw outputs, summary tables, figures, and deterministic seed derivation. The main benchmark uses 20 devices, 4 edge servers, 2 cloud servers, 40 tasks, population size 50, 150 iterations, and 30 paired scenario seeds.', size=12)
    set_paragraph(p[136], 'RIME and DBO isolate the parent mechanisms, TLBO-HHO [21] and CWTSSA represent enhanced population-based optimisers, and Greedy-ED is a lightweight coordinate-wise reference. Every method receives the same task and network scenario for a given run and is evaluated with the same metric code and fixed-reference reporting fitness.', size=12)
    set_paragraph(p[137], 'Population size and iteration count are equal for the population-based methods. Because RDHO-full includes a greedy seed and final local refinement, it uses more objective-function evaluations than the other population methods; Table 5 therefore reports NFE and runtime instead of implying a strictly NFE-matched comparison. Convergence curves show the fixed-reference incumbent before local refinement, while the final table reports the complete RDHO-full output.', size=12)
    set_paragraph(p[138], 'Table 2. System parameters used in the main experiment', size=10)
    set_paragraph(p[139], 'Table 3. Task-generation ranges and sampling probabilities', size=10)
    set_paragraph(p[140], 'Table 4. Algorithm and reproducibility parameters', size=10)

    # Dynamic result narrative.
    energy_best = best_name(main_raw, 'energy', False)
    delay_best = best_name(main_raw, 'delay', False)
    aoi_best = best_name(main_raw, 'aoi', False)
    qoe_best = best_name(main_raw, 'qoe', True)
    fairness_best = best_name(main_raw, 'fairness', True)
    csr_best = best_name(main_raw, 'csr', True)
    time_best = best_name(main_raw, 'runtime', False)

    set_paragraph(p[142], 'Fig. 2 reports the mean fixed-reference incumbent over the 150 population iterations. The curve excludes the final coordinate refinement, so it compares the population-search trajectories under one common reporting scale; RDHO-full\'s post-search gain is reported separately in Table 7.', size=12)
    set_paragraph(p[144], 'Fig. 2. Fixed-reference reporting-fitness convergence over 150 iterations; RDHO local refinement is excluded from the curve.', size=10)
    set_paragraph(p[145], 'Table 5. Main comparison over 30 paired runs (mean ± standard deviation); fitness is the fixed-reference reporting fitness.', size=10)
    set_paragraph(p[146], f'RDHO obtains the lowest mean reporting fitness ({rdho["fitness_mean"]:.4f} ± {rdho["fitness_std"]:.4f}); the relative reductions are {improvement_text}. The associated mean QoE, priority-aware fairness, and soft CSR are {rdho["qoe_mean"]:.4f}, {rdho["fairness_mean"]:.4f}, and {rdho["csr_mean"]:.4f}, respectively. These values support the overall weighted-objective claim but do not establish dominance on every raw metric.', size=12)
    set_paragraph(p[147], f'The metric-specific leaders are {energy_best} for energy, {delay_best} for delay, {aoi_best} for AoI, {qoe_best} for QoE, {fairness_best} for priority-aware fairness, {csr_best} for soft CSR, and {time_best} for runtime. RDHO-full also uses more NFE because of its final coordinate refinement, so its advantage should be read as solution-quality versus computational-cost trade-off.', size=12)
    set_paragraph(p[149], 'Fig. 3. Device-side energy proxy over 30 paired runs.', size=10)
    set_paragraph(p[150], f'Fig. 3 shows that {energy_best} has the lowest mean device-side energy proxy. RDHO optimises energy as one weighted component rather than as a sole objective.', size=12)
    set_paragraph(p[152], 'Fig. 4. Mean processing delay over 30 paired runs.', size=10)
    set_paragraph(p[153], f'Fig. 4 identifies {delay_best} as the lowest-delay method. RDHO\'s reporting-fitness advantage therefore cannot be attributed to latency alone.', size=12)
    set_paragraph(p[155], 'Fig. 5. Mean average-AoI approximation over 30 paired runs.', size=10)
    set_paragraph(p[156], f'Fig. 5 shows that {aoi_best} obtains the lowest mean AoI. The close relationship between delay and Eq. (8) is expected because average AoI contains the service-delay term.', size=12)
    set_paragraph(p[158], 'Fig. 6. Mean QoE proxy and user-level Jain fairness.', size=10)
    set_paragraph(p[159], f'{qoe_best} achieves the highest mean QoE proxy, while {fairness_best} achieves the highest mean user-level fairness. Fairness is computed after aggregating task QoE by source device and therefore measures equality of per-user mean service utility.', size=12)
    set_paragraph(p[161], 'Fig. 7. Soft QoS constraint-satisfaction ratio.', size=10)
    set_paragraph(p[162], f'{csr_best} reaches the highest mean soft CSR. A value below one means that some delay, battery-adjusted energy, or AoI threshold checks are violated; it does not indicate an illegal mode or computation-control value.', size=12)

    set_paragraph(p[164], 'A two-sided paired Wilcoxon signed-rank test compares RDHO with every baseline on the 30 paired reporting-fitness values. Holm correction controls the family-wise error rate; rank-biserial correlation and wins/ties/losses quantify effect magnitude and consistency.', size=12)
    set_paragraph(p[165], 'Table 6. Paired Wilcoxon tests for fixed-reference reporting fitness', size=10)
    sig_count = int((stats['p_holm'] < 0.05).sum()) if 'p_holm' in stats else 0
    set_paragraph(p[166], f'After Holm correction, {sig_count} of {len(stats)} comparisons are significant at 0.05. The effect sizes and run-level wins/ties/losses in Table 6 should be considered together with the mean differences; significance is not claimed for every raw metric.', size=12)

    set_paragraph(p[168], 'The ablation study separates the population-search contribution from the final coordinate refinement. RDHO-core disables only local refinement. Each remaining variant removes one named search component while retaining the same local-refinement setting as RDHO-full, eliminating the coupling present in the earlier implementation.', size=12)
    set_paragraph(p[169], 'Table 7. Isolated RDHO ablations over 30 paired runs', size=10)
    set_paragraph(p[171], 'Fig. 8. Fixed-reference fitness and soft CSR for RDHO ablations.', size=10)
    full = summary_row(ablation, 'algorithm', 'RDHO-full')
    core = summary_row(ablation, 'algorithm', 'RDHO-core')
    component = ablation[~ablation['algorithm'].isin(['RDHO-full', 'RDHO-core'])]
    comp_min = float(component['fitness_mean'].min())
    comp_max = float(component['fitness_mean'].max())
    set_paragraph(p[172], f'Local refinement reduces mean reporting fitness from {core["fitness_mean"]:.4f} for RDHO-core to {full["fitness_mean"]:.4f} for RDHO-full, a mean gain of {core["fitness_mean"]-full["fitness_mean"]:.4f}. The single-component variants span {comp_min:.4f}-{comp_max:.4f}. Because these component effects are comparatively small and a removed-component variant may occasionally match or exceed the full mean, the ablation supports the importance of local refinement more strongly than a claim that every search component is individually indispensable.', size=12)

    set_paragraph(p[174], 'The scalability study increases task count from 20 to 100 with the same device-edge-cloud counts and task-generation process. Ten paired runs are used at each size.', size=12)
    set_paragraph(p[175], 'Table 8. RDHO scalability under different task counts', size=10)
    set_paragraph(p[177], 'Fig. 9. RDHO reporting fitness, soft CSR, and runtime versus task count.', size=10)
    first_scale = scale.sort_values('task_number').iloc[0]
    last_scale = scale.sort_values('task_number').iloc[-1]
    set_paragraph(p[178], f'Mean reporting fitness changes from {first_scale["fitness_mean"]:.4f} at {int(first_scale["task_number"])} tasks to {last_scale["fitness_mean"]:.4f} at {int(last_scale["task_number"])} tasks, while soft CSR changes from {first_scale["csr_mean"]:.4f} to {last_scale["csr_mean"]:.4f}. Runtime grows from {first_scale["runtime_mean"]:.3f} s to {last_scale["runtime_mean"]:.3f} s. These results describe empirical scaling for the tested range rather than an asymptotic complexity proof.', size=12)

    set_paragraph(p[180], 'Two analyses vary the objective-weight vector and the dynamic search-penalty schedule. All final solutions are evaluated with the same fixed-reference reporting objective, so the reported fitness differences reflect changes in the solutions found rather than direct rescaling by λ(t).', size=12)
    set_paragraph(p[181], f'Across S1-S5, mean QoE ranges from {weight["qoe_mean"].min():.4f} to {weight["qoe_mean"].max():.4f}, priority-aware fairness from {weight["fairness_mean"].min():.4f} to {weight["fairness_mean"].max():.4f}, and soft CSR from {weight["csr_mean"].min():.4f} to {weight["csr_mean"].max():.4f}. Reporting fitness changes with the weight vector, as expected.', size=12)
    set_paragraph(p[182], 'S1 is the default QoE-fairness-preferred setting, S2 uses equal weights, S3 emphasises energy and delay, S4 further strengthens QoE and fairness, and S5 emphasises AoI. Because only RDHO is run in this analysis, it demonstrates within-method sensitivity and does not prove that the cross-algorithm ranking is invariant to weights.', size=12)
    set_paragraph(p[183], 'Table 9. RDHO objective-weight sensitivity over 30 paired scenarios', size=10)
    set_paragraph(p[185], 'Fig. 10. RDHO weight sensitivity for QoE, priority-aware fairness, and soft CSR.', size=10)
    set_paragraph(p[186], f'Across the nine λ₀-α combinations, mean soft CSR ranges from {penalty["csr_mean"].min():.4f} to {penalty["csr_mean"].max():.4f}, and fixed-reference reporting fitness ranges from {penalty["fitness_mean"].min():.4f} to {penalty["fitness_mean"].max():.4f}. The dynamic parameters influence search trajectories, while the fixed reporting coefficient remains one.', size=12)
    set_paragraph(p[187], 'Table 10. Dynamic-penalty sensitivity evaluated with fixed-reference reporting fitness', size=10)
    set_paragraph(p[190], 'Fig. 11. Dynamic-penalty sensitivity heatmaps for soft CSR and fixed-reference reporting fitness.', size=10)
    set_paragraph(p[191], 'The sensitivity results indicate whether RDHO is fragile within the tested parameter ranges. They do not establish optimal parameter values outside those ranges or universal robustness across untested task and network distributions.', size=12)

    set_paragraph(p[194], 'Fig. 12. Min-max normalised multi-metric comparison; energy, delay, and AoI are reversed so that larger radial values are better.', size=10)
    set_paragraph(p[195], 'For each metric, min-max normalisation is performed across the compared algorithms. Energy, delay, and AoI are transformed as 1-(x-min)/(max-min), while QoE and fairness use (x-min)/(max-min). The radar chart is descriptive; reporting fitness, soft CSR, runtime, and NFE remain in the tables.', size=12)
    set_paragraph(p[196], 'Overall, RDHO offers the lowest fixed-reference reporting fitness in the implemented suite, but the raw-metric leaders vary and RDHO-full incurs additional evaluations. The evidence therefore supports a strong weighted trade-off under the stated model, not universal dominance or an NFE-free improvement.', size=12)

    set_paragraph(p[198], 'This study investigated QoE- and user-fairness-aware task offloading in a heterogeneous cloud-edge-device MEC model. Each task selects a local, edge, or cloud execution mode and a bounded normalised computation-control value under fixed associations and rates. The objective combines device-side energy, processing delay, average AoI, a priority-weighted QoE proxy, and fairness over per-user mean QoE.', size=12)
    set_paragraph(p[199], f'In 30 paired simulations, RDHO achieves the lowest mean fixed-reference reporting fitness ({rdho["fitness_mean"]:.4f}) among the implemented baselines. Corrected paired tests, isolated ablations, scalability analysis, and parameter sensitivity provide complementary evidence. Local refinement produces a measurable solution-quality gain but also increases NFE, and several baselines remain preferable on individual raw metrics or runtime.', size=12)
    set_paragraph(p[200], 'The study remains limited by simulated offline task sets, fixed access and cloud associations, fixed communication rates, a load-sharing rather than explicit sum-capacity CPU model, a model-based QoE proxy without user-study calibration, and a restricted baseline suite. Future work should examine online arrivals, mobility and interference, explicit bandwidth/CPU allocation, NFE-matched comparisons, multiobjective Pareto methods, learning-based schedulers, and physical testbed validation.', size=12)
    set_paragraph(p[210], 'The simulation code, configuration files, tests, raw outputs, summary tables, and generated figures are available at https://github.com/Ryan-Yii/mec-rdho-offloading. The revised implementation separates search and reporting fitness, records NFE, and includes regenerated paired experiments. A versioned release or archival DOI should be created before submission so that the manuscript is tied to an immutable code snapshot. No proprietary, confidential, or human-subject data were used.', size=12)

    # Equations (1)-(26).
    tuple_formula = r'\tau_i=(L_i,C_i,D_i^{\max},A_i^{\max},E_i^{\mathrm{bud}},\Delta_i,b_i,\pi_i)'
    formulas = [
        r'f_i^{\mathrm{L}}=\frac{F_{m(i)}^{\mathrm{L}}(0.35+0.65r_i)}{1+0.07(n_{m(i)}^{\mathrm{L}}-1)}',
        r'f_i^{\mathrm{E}}=\frac{F_{g(i)}^{\mathrm{E}}(0.40+0.60r_i)}{1+0.035(n_{g(i)}^{\mathrm{E}}-1)}',
        r'f_i^{\mathrm{C}}=\frac{F_{c(i)}^{\mathrm{C}}(0.45+0.55r_i)}{1+0.020(n_{c(i)}^{\mathrm{C}}-1)}',
        r'D_i=\begin{cases}\frac{C_i}{f_i^{\mathrm{L}}},&z_i=0\\[2pt]\frac{L_i}{R_{m(i),g(i)}}+\frac{C_i}{f_i^{\mathrm{E}}}+\delta_{\mathrm{E}},&z_i=1\\[2pt]\frac{L_i}{R_{m(i),g(i)}}+\frac{L_i}{R_{g(i),c(i)}}+\frac{C_i}{f_i^{\mathrm{C}}}+\delta_{\mathrm{C}},&z_i=2\end{cases}',
        r'E_i^{\mathrm{L}}=\kappa_{m(i)}C_i\left(f_i^{\mathrm{L}}\right)^2',
        r'E_i^{\mathrm{E}}=p_{m(i)}\frac{L_i}{R_{m(i),g(i)}}+0.015\frac{C_i}{f_i^{\mathrm{E}}}',
        r'E_i^{\mathrm{C}}=p_{m(i)}\frac{L_i}{R_{m(i),g(i)}}+0.010\left(\frac{L_i}{R_{g(i),c(i)}}+\frac{C_i}{f_i^{\mathrm{C}}}\right)',
        r'A_i=\frac{1}{2}\Delta_i+D_i',
        r's_i^{D}=\exp\left(-\frac{D_i}{D_i^{\max}}\right)',
        r's_i^{E}=\exp\left(-\frac{E_i}{E_i^{\mathrm{bud}}}\right)',
        r's_i^{A}=\exp\left(-\frac{A_i}{A_i^{\max}}\right)',
        r'q_i=\pi_i\left(0.45s_i^{D}+0.30s_i^{E}+0.25s_i^{A}\right)',
        r'q_i\leftarrow\min\{1,\max\{0,q_i\}\}',
        r'\overline{q}_m=\frac{1}{|\mathcal{N}_m|}\sum_{i\in\mathcal{N}_m}q_i,\qquad J=\frac{\left(\sum_{m=1}^{M_{\mathrm{act}}}\overline{q}_m\right)^2}{M_{\mathrm{act}}\sum_{m=1}^{M_{\mathrm{act}}}\overline{q}_m^2}',
        r'\overline{E}=\frac{1}{N}\sum_i\frac{E_i}{E_i^{\mathrm{bud}}},\qquad \overline{D}=\frac{1}{N}\sum_i\frac{D_i}{D_i^{\max}},\qquad \overline{A}=\frac{1}{N}\sum_i\frac{A_i}{A_i^{\max}}',
        r'F_{\mathrm{base}}(X)=w_E\overline{E}+w_D\overline{D}+w_A\overline{A}+w_Q(1-\overline{q})+w_J(1-J)',
        r'\mathrm{CSR}(X)=\frac{1}{3N}\sum_i\left(\mathbb{I}[D_i\le D_i^{\max}]+\mathbb{I}[E_i\le b_iE_i^{\mathrm{bud}}]+\mathbb{I}[A_i\le A_i^{\max}]\right)',
        r'z_i\in\{0,1,2\}',
        r'0.2\le r_i\le 1.0',
        r'v(X)=1-\mathrm{CSR}(X)',
        r'F_{\mathrm{search}}(X,t)=F_{\mathrm{base}}(X)+\lambda(t)v(X)',
        r'\lambda(t)=\lambda_0\left(1+\frac{2t}{T_{\max}}\right)^{\alpha}',
        r'F_{\mathrm{report}}(X)=F_{\mathrm{base}}(X)+\lambda_{\mathrm{ref}}v(X)',
        r'\lambda_{\mathrm{ref}}=1',
        r'w(t)=0.5+0.3\cos\left(\pi\frac{t}{T_{\max}}\right)',
        r'X_i^{\mathrm{new}}=w(t)X_i^{\mathrm{RIME}}+\left(1-w(t)\right)X_i^{\mathrm{DBO}}',
    ]
    tuple_node, *nodes = build_math_nodes([tuple_formula, *formulas])
    set_display_equation(p[31], tuple_node)
    equation_paragraphs = [36,38,40,42,45,47,49,52,59,63,67,70,72,76,82,85,88,90,92,94,97,99,102,104,124,125]
    for number, (idx, node) in enumerate(zip(equation_paragraphs, nodes), start=1):
        set_equation(p[idx], node, number)

    # Tables.
    notation = [
        ['Symbol', 'Definition'],
        ['M, G, C, N', 'Device, edge-server, cloud-server, and task sets'],
        ['m(i), g(i), c(i)', 'Source device and fixed access-edge/cloud mappings for task i'],
        ['zᵢ', 'Execution mode: 0 local, 1 edge, 2 cloud'],
        ['rᵢ', 'Computational-computation-control value, 0.2≤rᵢ≤1'],
        ['Lᵢ, Cᵢ', 'Input bits and required CPU cycles'],
        ['Dᵢ, Eᵢ, Aᵢ', 'Delay, device-side energy proxy, and average AoI'],
        ['Dᵢᵐᵃˣ, Eᵢᵇᵘᵈ, Aᵢᵐᵃˣ', 'Task-specific delay, nominal energy, and AoI thresholds'],
        ['bᵢ, πᵢ, Δᵢ', 'Battery ratio, task priority, and update interval'],
        ['fᵢᴸ, fᵢᴱ, fᵢᶜ', 'Load-adjusted effective execution frequencies'],
        ['qᵢ, q̄ₘ', 'Task QoE proxy and per-user mean QoE'],
        ['J', 'Jain fairness over active-user mean QoE'],
        ['CSR', 'Soft delay/energy/AoI satisfaction ratio'],
        ['Fbase', 'Continuous weighted base objective'],
        ['Fsearch', 'Iteration-dependent internal search fitness'],
        ['Freport', 'Fixed-reference cross-algorithm reporting fitness'],
        ['λ(t), λref', 'Dynamic search coefficient and fixed reporting coefficient'],
        ['NFE', 'Number of objective-function evaluations'],
    ]
    set_table_data(doc.tables[0], notation, widths=[1.65,4.75], font_size=9)

    algorithm_rows = [
        ['Require:', 'Task/network instance; weights; λ0, α; population P; iterations T'],
        ['Ensure:', 'Best solution X* evaluated by reporting fitness and its raw metrics'],
        ['1', 'Generate half Gaussian and half uniform individuals'],
        ['2', 'Insert one greedy coordinate seed and up to three perturbations'],
        ['3', 'Evaluate base metrics and initialise search/reporting incumbents'],
        ['4', 'for t=1,…,T do'],
        ['5', 'Recompute parent search fitness with the current λ(t)'],
        ['6', 'Sort population; assign adaptive roles; preserve top 10% elites'],
        ['7', 'Update producers with fused RIME-DBO operators'],
        ['8', 'Update followers by puncture or foraging'],
        ['9', 'Update scouts by theft or decaying Cauchy mutation'],
        ['10', 'Round/clip z and clip r to their legal bounds'],
        ['11', 'Evaluate candidates once to obtain the base objective and CSR'],
        ['12', 'Derive candidate and parent search fitness using the same λ(t)'],
        ['13', 'Greedily accept candidates with lower current search fitness'],
        ['14', 'Update the independent reporting-fitness incumbent'],
        ['15', 'end for'],
        ['16', 'If enabled, run two coordinate-refinement passes'],
        ['17', 'Return X*, reporting fitness, raw metrics, and NFE'],
    ]
    set_table_data(doc.tables[1], algorithm_rows, widths=[0.75,5.65], font_size=9)

    table2 = [
        ['Parameter', 'Value'],
        ['Devices / edge / cloud servers', '20 / 4 / 2'],
        ['Task count', '40 heterogeneous tasks'],
        ['Device / edge / cloud CPU', '0.8-2.2 / 8-18 / 25-40 GHz'],
        ['Device transmit power', '0.2-0.8 W'],
        ['Device-edge / edge-cloud rate', '8-30 / 60-150 Mbps'],
        ['Device energy coefficient', '(0.8-1.4)×10⁻²⁷'],
        ['Fixed edge / cloud overhead', '0.010 / 0.055 s'],
    ]
    set_table_data(doc.tables[2], table2, widths=[3.0,3.4], font_size=9)

    table3 = [
        ['Parameter', 'Compute-intensive', 'Data-intensive', 'Real-time', 'Lightweight'],
        ['Sampling probability', '0.28', '0.24', '0.28', '0.20'],
        ['Input data (MB)', '8-35', '30-90', '2-15', '0.5-5'],
        ['CPU cycles (Gcycles)', '1.8-4.5', '0.8-2.2', '0.5-1.8', '0.1-0.8'],
        ['Maximum delay (s)', '1.5-3.0', '2.5-5.0', '0.35-1.0', '0.8-2.0'],
        ['AoI threshold (s)', '2.0-3.0', '3.0-5.0', '0.5-1.0', '1.0-2.0'],
        ['Energy budget (J)', '2.5-6.0', '2.0-5.0', '1.0-3.5', '0.5-2.0'],
        ['Battery ratio', '0.45-1.0', '0.40-0.95', '0.55-1.0', '0.50-1.0'],
        ['Priority / update interval (s)', '0.60-0.90 / 0.50-1.00', '0.50-0.80 / 0.80-1.50', '0.80-1.00 / 0.10-0.30', '0.40-0.70 / 0.30-0.60'],
    ]
    set_table_data(doc.tables[3], table3, widths=[1.45,1.25,1.25,1.25,1.25], font_size=8.2)

    table4 = [
        ['Parameter', 'Value'],
        ['Population / iterations', '50 / 150'],
        ['Default weights (E,D,A,Q,J)', '0.15, 0.15, 0.20, 0.25, 0.25'],
        ['Dynamic penalty (λ0, α)', '1.0, 2.0'],
        ['Fixed reporting coefficient', '1.0'],
        ['Initialisation / elite ratio', '50% Gaussian + 50% uniform / 10%'],
        ['Local-refinement grid', 'modes {0,1,2}; ratios {0.25,0.40,0.60,0.80,1.00}; ≤2 passes'],
        ['Independent paired runs', '30 main and ablation; 10 per scalability size'],
        ['Compared methods', 'RDHO, RIME, DBO, TLBO-HHO, CWTSSA, Greedy-ED'],
    ]
    set_table_data(doc.tables[4], table4, widths=[2.25,4.15], font_size=9)

    main_table = [['Alg.', 'Reporting fitness', 'Energy (J)', 'Delay (s)', 'AoI (s)', 'QoE', 'Priority-aware fairness', 'Soft CSR', 'Time (s)', 'NFE']]
    for name in ['RDHO', 'RIME', 'DBO', 'TLBO-HHO', 'CWTSSA', 'Greedy-ED']:
        row = summary_row(main_summary, 'algorithm', name)
        main_table.append([
            name,
            fmt(row['fitness_mean'], row['fitness_std'], 3),
            fmt(row['energy_mean'], row['energy_std'], 1),
            fmt(row['delay_mean'], row['delay_std'], 3),
            fmt(row['aoi_mean'], row['aoi_std'], 3),
            fmt(row['qoe_mean'], row['qoe_std'], 3),
            fmt(row['fairness_mean'], row['fairness_std'], 3),
            fmt(row['csr_mean'], row['csr_std'], 3),
            fmt(row['runtime_mean'], row['runtime_std'], 3),
            f'{row["nfe_mean"]:.0f}',
        ])
    set_table_data(doc.tables[5], main_table, widths=[0.58,0.9,0.75,0.7,0.7,0.58,0.78,0.68,0.68,0.55], font_size=7.4)

    stats_table = [['Comparison', 'Raw p', 'Holm p', 'Rank-\nbiserial', 'W/T/L']]
    for _, row in stats.iterrows():
        stats_table.append([
            row['comparison'],
            f'{row["p_value"]:.2e}',
            f'{row["p_holm"]:.2e}',
            f'{row["rank_biserial"]:.3f}',
            f'{int(row["wins"])}/{int(row["ties"])}/{int(row["losses"])}',
        ])
    set_table_data(doc.tables[6], stats_table, widths=[1.50,0.85,0.85,1.00,0.80], font_size=8.2)

    ablation_table = [['Variant', 'Reporting fitness', 'QoE', 'Priority-aware fairness', 'Soft CSR', 'Time (s)']]
    for name in ['RDHO-full', 'RDHO-core', 'RDHO-w/o dual-source initialization', 'RDHO-w/o adaptive role allocation', 'RDHO-w/o elite preservation', 'RDHO-w/o dynamic penalty']:
        row = summary_row(ablation, 'algorithm', name)
        short = ablation_label(name)
        ablation_table.append([
            short,
            fmt(row['fitness_mean'], row['fitness_std'], 3),
            fmt(row['qoe_mean'], row['qoe_std'], 3),
            fmt(row['fairness_mean'], row['fairness_std'], 3),
            fmt(row['csr_mean'], row['csr_std'], 3),
            fmt(row['runtime_mean'], row['runtime_std'], 3),
        ])
    set_table_data(doc.tables[7], ablation_table, widths=[1.45,1.0,0.7,0.85,0.75,0.75], font_size=8.1)

    scale_table = [['Tasks', 'Reporting fitness', 'Soft CSR', 'Time (s)', 'NFE']]
    for _, row in scale.sort_values('task_number').iterrows():
        scale_table.append([
            f'{int(row["task_number"])}',
            fmt(row['fitness_mean'], row['fitness_std'], 3),
            fmt(row['csr_mean'], row['csr_std'], 3),
            fmt(row['runtime_mean'], row['runtime_std'], 3),
            f'{row.get("nfe_mean", np.nan):.0f}' if 'nfe_mean' in row and not pd.isna(row.get('nfe_mean')) else '—',
        ])
    set_table_data(doc.tables[8], scale_table, widths=[0.75,1.35,1.0,1.0,0.8], font_size=8.8)

    weight_table = [['Set', 'Weights (E,D,A,Q,J)', 'Reporting fitness', 'QoE', 'Priority-aware fairness', 'Soft CSR', 'Time (s)']]
    for setting in ['S1','S2','S3','S4','S5']:
        row = summary_row(weight, 'setting', setting)
        weight_table.append([
            setting, str(row['weights']), fmt(row['fitness_mean'], row['fitness_std'], 3), fmt(row['qoe_mean'], row['qoe_std'], 3),
            fmt(row['fairness_mean'], row['fairness_std'], 3), fmt(row['csr_mean'], row['csr_std'], 3), fmt(row['runtime_mean'], row['runtime_std'], 3),
        ])
    set_table_data(doc.tables[9], weight_table, widths=[0.45,1.5,1.0,0.7,0.85,0.75,0.75], font_size=7.8)

    penalty_table = [['λ₀', 'α', 'Reporting fitness', 'QoE', 'Priority-aware fairness', 'Soft CSR', 'Time (s)']]
    for _, row in penalty.sort_values(['lambda0','alpha']).iterrows():
        penalty_table.append([
            f'{row["lambda0"]:.1f}', f'{row["alpha"]:.1f}', fmt(row['fitness_mean'], row['fitness_std'], 3),
            fmt(row['qoe_mean'], row['qoe_std'], 3), fmt(row['fairness_mean'], row['fairness_std'], 3),
            fmt(row['csr_mean'], row['csr_std'], 3), fmt(row['runtime_mean'], row['runtime_std'], 3),
        ])
    set_table_data(doc.tables[10], penalty_table, widths=[0.55,0.55,1.15,0.75,0.9,0.8,0.8], font_size=8.1)

    # Replace result figures; preserve original architecture figure.
    figure_map = {
        143: (REPO / 'results/figures/convergence_curve.png', 4.85),
        148: (REPO / 'results/figures/energy_comparison.png', 4.85),
        151: (REPO / 'results/figures/delay_comparison.png', 4.85),
        154: (REPO / 'results/figures/aoi_comparison.png', 4.85),
        157: (REPO / 'results/figures/qoe_fairness_comparison.png', 4.85),
        160: (REPO / 'results/figures/csr_comparison.png', 4.85),
        170: (REPO / 'results/figures/ablation_study_multicolor.png', 4.85),
        176: (REPO / 'results/figures/scalability.png', 4.85),
        184: (REPO / 'results/sensitivity/figures/weight_sensitivity_qoe_fairness_csr.png', 4.85),
        189: (REPO / 'results/sensitivity/figures/penalty_sensitivity_heatmaps.png', 4.85),
        193: (REPO / 'results/figures/radar_chart.png', 3.82),
    }
    for idx, (path, width) in figure_map.items():
        if not path.exists():
            raise FileNotFoundError(path)
        replace_picture(p[idx], path, width)

    # Keep all revised body text at 12 pt and use consistent paragraph spacing.
    for paragraph in doc.paragraphs:
        if paragraph.style.name in {'Body Text', 'First Paragraph', 'Normal', 'Abstract'}:
            paragraph.paragraph_format.space_after = Pt(3)
            paragraph.paragraph_format.line_spacing = 1.05

    doc.core_properties.title = 'RIME-DBO-Based QoE- and Fairness-Aware Task Offloading in Mobile Edge Computing'
    doc.core_properties.subject = 'Marked SCI revision aligned with reproducibility repository'
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == '__main__':
    main()
